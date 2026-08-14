"""
Export pipeline — renders a LabGuide to multiple output formats.

Supported:
  - Markdown (.md)
  - PDF (via Chrome headless)
  - DOCX (via python-docx)
  - HTML (standalone, embedded images)
  - MkDocs site (generate docs/ tree + mkdocs.yml, optional git push)
  - Narrated video (TTS + ffmpeg overlay)
"""
from __future__ import annotations
import re
import subprocess
from pathlib import Path
from datetime import datetime
import html as _html_mod
from html.parser import HTMLParser as _HTMLParser

from lab_guide_automator.models import LabGuide, LabSection, LabStep


def _md_to_html(text: str) -> str:
    """Convert simple Markdown to HTML (headings, bold, italic, lists, hr, paragraphs)."""
    lines = text.split("\n")
    out = []
    in_ul = False
    in_ol = False

    def close_lists():
        nonlocal in_ul, in_ol
        if in_ul:
            out.append("</ul>")
            in_ul = False
        if in_ol:
            out.append("</ol>")
            in_ol = False

    def inline(s: str) -> str:
        s = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', s)
        s = re.sub(r'\*(.+?)\*', r'<em>\1</em>', s)
        s = re.sub(r'\[(.+?)\]\(([^\)]+)\)', r'<a href="\2" target="_blank" rel="noopener noreferrer">\1</a>', s)
        s = re.sub(r'`(.+?)`', r'<code>\1</code>', s)
        return s

    for line in lines:
        # Headings
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            close_lists()
            lvl = len(m.group(1)) + 1  # shift: # → h2, ## → h3, etc.
            lvl = min(lvl, 6)
            out.append(f"<h{lvl}>{inline(m.group(2))}</h{lvl}>")
            continue
        # HR
        if re.match(r'^-{3,}$|^\*{3,}$', line.strip()):
            close_lists()
            out.append("<hr>")
            continue
        # Unordered list
        m = re.match(r'^[-*]\s+(.*)', line)
        if m:
            if not in_ul:
                close_lists()
                out.append("<ul>")
                in_ul = True
            out.append(f"<li>{inline(m.group(1))}</li>")
            continue
        # Ordered list
        m = re.match(r'^\d+\.\s+(.*)', line)
        if m:
            if not in_ol:
                close_lists()
                out.append("<ol>")
                in_ol = True
            out.append(f"<li>{inline(m.group(1))}</li>")
            continue
        # Empty line
        if not line.strip():
            close_lists()
            continue
        # Paragraph
        close_lists()
        out.append(f"<p>{inline(line)}</p>")

    close_lists()
    return "\n".join(out)


def _add_link_targets(html: str) -> str:
    """Inject target="_blank" rel="noopener noreferrer" into every <a> tag that lacks it."""
    def _patch(m: re.Match) -> str:
        tag = m.group(0)
        if 'target=' not in tag:
            tag = tag.rstrip('>').rstrip('/') + ' target="_blank" rel="noopener noreferrer">'
        return tag
    return re.sub(r'<a\s[^>]*>', _patch, html, flags=re.IGNORECASE)


# ---------------------------------------------------------------------------
# HTML → clean Markdown helper (handles Quill output)
# ---------------------------------------------------------------------------

class _QuillToMd(_HTMLParser):
    """Lightweight Quill HTML → Markdown converter (stdlib only)."""

    _BLOCK_TAGS = {"p", "div", "li", "h1", "h2", "h3", "h4", "h5", "h6",
                   "blockquote"}
    _HEADING_MAP = {"h1": "#", "h2": "##", "h3": "###", "h4": "####", "h5": "#####", "h6": "######"}
    _IGNORE_TAGS = {"span", "em", "i", "u", "s", "code", "thead", "tbody", "table"}

    def __init__(self):
        super().__init__()
        self._parts: list[str] = []
        self._bold = 0
        self._italic = 0
        self._link_href = ""
        self._in_li = False
        self._li_prefix = ""
        self._in_pre = False
        self._pre_buf: list[str] = []
        # Table state
        self._in_table = False
        self._in_th = False
        self._in_td = False
        self._col_idx = 0
        self._row_idx = 0
        self._table_row_buf: list[str] = []
        self._table_header_done = False

    def handle_starttag(self, tag, attrs):
        attr = dict(attrs)
        if tag == "strong" or tag == "b":
            self._bold += 1
        elif tag == "em" or tag == "i":
            self._italic += 1
        elif tag == "a":
            self._link_href = attr.get("href", "")
        elif tag in ("ol", "ul"):
            self._li_prefix = "1. " if tag == "ol" else "- "
        elif tag == "li":
            self._in_li = True
            self._parts.append(self._li_prefix)
        elif tag in ("br",):
            self._parts.append("  \n")
        elif tag == "pre":
            self._in_pre = True
            self._pre_buf = []
        elif tag == "code" and not self._in_pre:
            pass  # inline code — data flows through normally
        elif tag == "hr":
            self._parts.append("\n\n---\n\n")
        elif tag == "table":
            self._in_table = True
            self._row_idx = 0
            self._table_header_done = False
            self._parts.append("\n\n")
        elif tag == "tr":
            self._table_row_buf = []
            self._col_idx = 0
        elif tag == "th":
            self._in_th = True
            self._col_idx += 1
        elif tag == "td":
            self._in_td = True
            self._col_idx += 1
        elif tag in self._HEADING_MAP:
            self._parts.append(self._HEADING_MAP[tag] + " ")

    def handle_endtag(self, tag):
        if tag == "strong" or tag == "b":
            self._bold = max(0, self._bold - 1)
        elif tag == "em" or tag == "i":
            self._italic = max(0, self._italic - 1)
        elif tag == "a":
            self._link_href = ""
        elif tag in self._BLOCK_TAGS or tag in self._HEADING_MAP:
            self._parts.append("\n\n")
            self._in_li = False
        elif tag == "pre":
            content = "".join(self._pre_buf).strip()
            self._parts.append(f"\n\n```\n{content}\n```\n\n")
            self._in_pre = False
            self._pre_buf = []
        elif tag in ("th", "td"):
            cell = "".join(self._table_row_buf).strip()
            self._parts.append(f"| {cell} ")
            self._table_row_buf = []
            self._in_th = False
            self._in_td = False
        elif tag == "tr":
            self._parts.append("|\n")
            self._row_idx += 1
            # After the first row (header), emit the separator
            if self._row_idx == 1 and not self._table_header_done:
                # Count columns from pipe chars in last row
                last_row = "".join(self._parts).split("\n")[-2]
                col_count = last_row.count("|") - 1
                if col_count < 1:
                    col_count = 1
                self._parts.append("|" + " --- |" * col_count + "\n")
                self._table_header_done = True
        elif tag == "table":
            self._in_table = False
            self._parts.append("\n\n")

    def handle_data(self, data):
        # Strip the drag-handle glyph that may have leaked in
        data = data.replace("\u283f", "").replace("\u28ff", "")
        if not data:
            return
        if self._bold:
            data = f"**{data}**"
        if self._italic:
            data = f"*{data}*"
        if self._link_href:
            data = f"[{data}]({self._link_href})"
        # Route table cell data to buffer, not main parts
        if self._in_th or self._in_td:
            self._table_row_buf.append(data)
        elif self._in_pre:
            self._pre_buf.append(data)
        else:
            self._parts.append(data)

    def handle_entityref(self, name):
        self._parts.append(_html_mod.unescape(f"&{name};"))

    def handle_charref(self, name):
        self._parts.append(_html_mod.unescape(f"&#{name};"))

    def result(self) -> str:
        import re
        text = "".join(self._parts)
        # Collapse excessive blank lines, strip leading/trailing whitespace
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Replace non-breaking spaces
        text = text.replace("\xa0", " ").replace("&nbsp;", " ")
        return text.strip()


def _html_to_md(html: str) -> str:
    """Convert Quill-generated HTML to clean Markdown text."""
    if not html:
        return ""
    # Strip the drag-handle character at raw level too
    html = html.replace("\u283f", "").replace("⠿", "")
    parser = _QuillToMd()
    parser.feed(html)
    return parser.result()


def _md_strip(text: str) -> str:
    """Strip Markdown syntax from text → plain text suitable for DOCX paragraphs."""
    if not text:
        return ""
    import re as _re
    # Remove heading markers
    text = _re.sub(r'^#{1,6}\s+', '', text, flags=_re.MULTILINE)
    # Bold/italic
    text = _re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = _re.sub(r'\*(.+?)\*', r'\1', text)
    text = _re.sub(r'__(.+?)__', r'\1', text)
    text = _re.sub(r'_(.+?)_', r'\1', text)
    # Links
    text = _re.sub(r'\[(.+?)\]\([^\)]+\)', r'\1', text)
    # Inline code
    text = _re.sub(r'`(.+?)`', r'\1', text)
    # HR
    text = _re.sub(r'^[-*]{3,}$', '', text, flags=_re.MULTILINE)
    # List bullets
    text = _re.sub(r'^[-*]\s+', '', text, flags=_re.MULTILINE)
    text = _re.sub(r'^\d+\.\s+', '', text, flags=_re.MULTILINE)
    return text.strip()


def _to_plain(text: str) -> str:
    """Convert either Quill HTML or Markdown to clean plain text for DOCX paragraphs."""
    if not text or not text.strip():
        return ""
    if text.strip().startswith("<"):
        # HTML → markdown → strip markdown markers → plain text
        return _md_strip(_html_to_md(text))
    return _md_strip(text)


def _ffmpeg_bin() -> str:
    for candidate in ["ffmpeg", "/opt/homebrew/bin/ffmpeg", "/usr/local/bin/ffmpeg"]:
        try:
            subprocess.run([candidate, "-version"], capture_output=True, check=True)
            return candidate
        except (FileNotFoundError, subprocess.CalledProcessError):
            continue
    return "ffmpeg"


# ─────────────────────────────────────────────────────────────
# Markdown renderer (shared base for all text exports)
# ─────────────────────────────────────────────────────────────

_CALLOUT_MD = {
    "expected_result": (">", "**✓ Expected Result:**"),
    "note":            (">", "**📝 Note:**"),
    "caution":         (">", "**⚠ Caution:**"),
    "congratulations": (">", "**🎉 Congratulations:**"),
}


def _render_block_md(blk, include_screenshots: bool = True) -> list[str]:
    """Return markdown lines for a single ContentBlock."""
    lines: list[str] = []
    if blk.type == "text" and blk.content:
        lines += [blk.content, ""]
    elif blk.type == "screenshot" and blk.path and include_screenshots:
        fname = Path(blk.path).name
        lines += [f"![]({fname})", ""]
    elif blk.type == "callout":
        prefix, label = _CALLOUT_MD.get(blk.callout_type, (">", f"**{blk.callout_type}:**"))
        body = blk.content or ""
        lines += [f"{prefix} {label} {body}", ""]
    return lines


def render_markdown(guide: LabGuide, include_screenshots: bool = True) -> str:
    lines: list[str] = []
    m = guide.metadata

    lines += [
        f"# {m.title}",
        "",
        f"**Version:** {m.version}  |  "
        f"**Author:** {m.author or 'Unknown'}  |  "
        f"**Date:** {m.date}  |  "
        f"**Difficulty:** {m.difficulty.capitalize()}  |  "
        f"**Duration:** {m.lab_duration_minutes} min",
        "",
    ]

    if m.tags:
        lines += [f"**Tags:** {', '.join(m.tags)}", ""]
    if m.prerequisites:
        lines += ["## Prerequisites", ""]
        for p in m.prerequisites:
            lines += [f"- {p}"]
        lines += [""]

    if guide.introduction:
        lines += ["## Introduction", "", guide.introduction, ""]

    if guide.learning_objectives:
        lines += ["## Learning Objectives", ""]
        for obj in guide.learning_objectives:
            lines += [f"- {obj.text}"]
        lines += [""]

    for sec in guide.sections:
        lines += [f"## {sec.title}", ""]
        if sec.overview:
            lines += [sec.overview, ""]
        # Section blocks (new system)
        sec_blocks = getattr(sec, "blocks", []) or []
        if sec_blocks:
            for blk in sec_blocks:
                lines += _render_block_md(blk, include_screenshots)
        elif include_screenshots:
            for ss in getattr(sec, "screenshots", []):
                lines += [f"![]({ss.path})", ""]

        for step in sec.steps:
            lines += [f"### Step {step.order}: {step.title}", ""]
            lines += [_html_to_md(step.instruction), ""]
            if step.code_blocks:
                for cb in step.code_blocks:
                    lines += ["```", cb, "```", ""]
            # Step blocks (new system)
            step_blocks = getattr(step, "blocks", []) or []
            if step_blocks:
                for blk in step_blocks:
                    lines += _render_block_md(blk, include_screenshots)
            elif include_screenshots and step.screenshots:
                for ss in step.screenshots:
                    lines += [f"![]({ss.path})", ""]
            # Legacy dedicated fields
            if getattr(step, "expected_result", ""):
                lines += [f"> **✓ Expected Result:** {step.expected_result}", ""]
            if getattr(step, "notes", ""):
                lines += [f"> **📝 Note:** {step.notes}", ""]

    if guide.conclusion:
        lines += ["## Conclusion", "", guide.conclusion, ""]

    return "\n".join(lines)


def export_markdown(guide: LabGuide, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(render_markdown(guide))
    return output_path


# ─────────────────────────────────────────────────────────────
# HTML export (Moodle-compatible)
# ─────────────────────────────────────────────────────────────

_HTML_CSS = """
body { font-family: 'Segoe UI', Arial, sans-serif; max-width: 900px; margin: 0 auto; padding: 2rem; color: #1a1a1a; }
h1 { color: #005073; border-bottom: 3px solid #00bceb; padding-bottom: .5rem; }
h2 { color: #005073; margin-top: 2rem; }
h3 { color: #1f7a8c; }
.meta { color: #555; font-size: .9rem; margin-bottom: 1.5rem; }
.tag { background: #e8f7fc; color: #005073; border-radius: 4px; padding: 2px 8px; font-size: .8rem; }
.objective li { margin: .4rem 0; }
.step { border-left: 4px solid #00bceb; padding: 0 1rem; margin: 1rem 0; }
.callout { padding: .5rem 1rem; border-radius: 0 4px 4px 0; margin-top: .6rem; font-size: .9rem; border-left-width: 4px; border-left-style: solid; }
.callout-expected { background: #f0faf0; border-left-color: #4caf50; }
.callout-note     { background: #e3f2fd; border-left-color: #2196f3; }
.callout-caution  { background: #fff3e0; border-left-color: #ff9800; }
.callout-congrats { background: #f3e5f5; border-left-color: #9c27b0; }
.callout-tip      { background: #e0f2f1; border-left-color: #009688; }
code, pre { background: #f4f4f4; border-radius: 4px; padding: .2rem .4rem; font-family: monospace; }
pre { padding: 1rem; overflow-x: auto; }
img { max-width: 100%; border: 1px solid #ddd; border-radius: 4px; margin: .5rem 0; display: block; }
p { margin: .6rem 0; line-height: 1.6; }
"""

_CALLOUT_HTML = {
    "expected_result": ("callout-expected", "✓ Expected Result"),
    "note":            ("callout-note",     "📝 Note"),
    "caution":         ("callout-caution",  "⚠️ Caution"),
    "congratulations": ("callout-congrats", "🎉 Congratulations"),
    "tip":             ("callout-tip",      "💡 Tip"),
}

# Base screenshots directory — resolved relative to this file's project root
_SCREENSHOTS_DIR = Path(__file__).parent.parent / "data" / "screenshots"


def _resolve_screenshot_path(blk_path: str) -> Path:
    """Resolve a block screenshot path (bare filename or screenshots/filename) to an absolute Path."""
    p = Path(blk_path)
    if p.is_absolute() and p.exists():
        return p
    # strip leading 'screenshots/' prefix if present
    name = p.name
    candidate = _SCREENSHOTS_DIR / name
    if candidate.exists():
        return candidate
    return p


def _render_block_html(blk, embed: bool = False) -> str:
    import base64
    if blk.type == "text" and blk.content:
        # content is already Quill-generated HTML — return as-is, no extra <p> wrap
        return _add_link_targets(blk.content)
    elif blk.type == "screenshot" and blk.path:
        p = _resolve_screenshot_path(blk.path)
        if embed and p.exists():
            b64 = base64.b64encode(p.read_bytes()).decode()
            src = f"data:image/{p.suffix.lstrip('.')};base64,{b64}"
        elif p.exists():
            src = str(p)
        else:
            src = blk.path
        return f'<img src="{src}" alt="">'
    elif blk.type == "callout":
        css, label = _CALLOUT_HTML.get(blk.callout_type, ("callout-note", blk.callout_type or "Note"))
        body = blk.content or ""
        return f'<div class="callout {css}"><strong>{label}</strong>{("<br>" + body) if body else ""}</div>'
    return ""


def render_html(guide: LabGuide, embed_screenshots: bool = False) -> str:
    import base64
    m = guide.metadata
    tags_html = " ".join(f'<span class="tag">{t}</span>' for t in m.tags)

    parts = [f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{m.title}</title>
<style>{_HTML_CSS}</style>
</head>
<body>
<h1>{m.title}</h1>
<div class="meta">
  Version {m.version} &nbsp;|&nbsp; {m.author or 'Unknown'} &nbsp;|&nbsp; {m.date}
  &nbsp;|&nbsp; {m.difficulty.capitalize()} &nbsp;|&nbsp; {m.lab_duration_minutes} min
  <br>{tags_html}
</div>"""]

    if m.prerequisites:
        items = "".join(f"<li>{p}</li>" for p in m.prerequisites)
        parts.append(f"<h2>Prerequisites</h2><ul>{items}</ul>")

    if guide.introduction:
        parts.append(f"<h2>Introduction</h2>{_md_to_html(guide.introduction)}")

    if guide.learning_objectives:
        items = "".join(f"<li>{o.text}</li>" for o in guide.learning_objectives)
        parts.append(f'<h2>Learning Objectives</h2><ul class="objective">{items}</ul>')

    for sec in guide.sections:
        parts.append(f"<h2>{sec.title}</h2>")
        if sec.overview:
            parts.append(f"<p>{sec.overview}</p>")

        # Section blocks (new system)
        sec_blocks = getattr(sec, "blocks", []) or []
        if sec_blocks:
            for blk in sec_blocks:
                h = _render_block_html(blk, embed=embed_screenshots)
                if h:
                    parts.append(h)
        else:
            for ss in getattr(sec, "screenshots", []):
                img_path = _resolve_screenshot_path(ss.path)
                if embed_screenshots and img_path.exists():
                    import base64 as _b64
                    src = (f"data:image/{img_path.suffix.lstrip('.')};base64,"
                           + _b64.b64encode(img_path.read_bytes()).decode())
                else:
                    src = str(img_path) if img_path.exists() else ss.path
                parts.append(f'<img src="{src}" alt="">')

        for step in sec.steps:
            parts.append(f'<div class="step">')
            parts.append(f"<h3>Step {step.order}: {step.title}</h3>")
            if step.instruction and step.instruction.strip():
                instr = step.instruction.strip()
                # Already HTML from Quill — output raw; plain text → convert markdown
                if instr.startswith("<"):
                    parts.append(_add_link_targets(instr))
                else:
                    parts.append(_md_to_html(instr))
            for cb in step.code_blocks:
                parts.append(f"<pre><code>{cb}</code></pre>")

            # Step blocks (new system)
            step_blocks = getattr(step, "blocks", []) or []
            if step_blocks:
                for blk in step_blocks:
                    h = _render_block_html(blk, embed=embed_screenshots)
                    if h:
                        parts.append(h)
            else:
                for ss in step.screenshots:
                    img_path = _resolve_screenshot_path(ss.path)
                    if embed_screenshots and img_path.exists():
                        import base64 as _b64
                        src = (f"data:image/{img_path.suffix.lstrip('.')};base64,"
                               + _b64.b64encode(img_path.read_bytes()).decode())
                    else:
                        src = str(img_path) if img_path.exists() else ss.path
                    parts.append(f'<img src="{src}" alt="">')
            # Legacy dedicated fields
            if getattr(step, "expected_result", ""):
                parts.append(f'<div class="callout callout-expected"><strong>✓ Expected Result</strong><br>{step.expected_result}</div>')
            if getattr(step, "notes", ""):
                parts.append(f'<div class="callout callout-note"><strong>📝 Note</strong><br>{step.notes}</div>')

            parts.append("</div>")

    if guide.conclusion:
        parts.append(f"<h2>Conclusion</h2>{_md_to_html(guide.conclusion)}")

    parts.append("</body></html>")
    return "\n".join(parts)


def export_html(guide: LabGuide, output_path: Path, embed_screenshots: bool = False) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(render_html(guide, embed_screenshots=embed_screenshots))
    return output_path


# ─────────────────────────────────────────────────────────────
# PDF export (weasyprint)
# ─────────────────────────────────────────────────────────────

# ─────────────────────────────────────────────────────────────
# PDF export (Chrome headless)
# ─────────────────────────────────────────────────────────────

def export_pdf(guide: LabGuide, output_path: Path) -> Path:
    """Generate PDF via Chrome headless --print-to-pdf (no native lib deps)."""
    import tempfile, os

    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Write a self-contained HTML with embedded images
    with tempfile.NamedTemporaryFile(suffix=".html", delete=False, mode="w", encoding="utf-8") as f:
        f.write(render_html(guide, embed_screenshots=True))
        tmp_html = f.name

    chrome_candidates = [
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Chromium.app/Contents/MacOS/Chromium",
        "/usr/bin/google-chrome",
        "/usr/bin/chromium-browser",
        "/usr/bin/chromium",
    ]
    chrome = next((c for c in chrome_candidates if Path(c).exists()), None)
    if not chrome:
        os.unlink(tmp_html)
        raise RuntimeError(
            "Chrome/Chromium not found. Install Google Chrome or Chromium to export PDF."
        )

    try:
        result = subprocess.run(
            [
                chrome,
                "--headless=new",
                "--disable-gpu",
                "--no-sandbox",
                "--disable-dev-shm-usage",
                f"--print-to-pdf={output_path}",
                "--print-to-pdf-no-header",
                f"file://{tmp_html}",
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(f"Chrome PDF export failed: {result.stderr[:500]}")
    finally:
        os.unlink(tmp_html)

    return output_path


# ─────────────────────────────────────────────────────────────
# DOCX export (python-docx)
# ─────────────────────────────────────────────────────────────

def export_docx(guide: LabGuide, output_path: Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import html as _html_mod2
        from html.parser import HTMLParser
    except ImportError:
        raise RuntimeError("python-docx is required: pip install python-docx")

    _CALLOUT_LABEL = {
        "expected_result": "✓ Expected Result",
        "note":            "📝 Note",
        "caution":         "⚠️ Caution",
        "congratulations": "🎉 Congratulations",
        "tip":             "💡 Tip",
        "team_challenge":  "🏆 Team Challenge",
    }

    _CALLOUT_COLOR = {
        "expected_result": RGBColor(0xE6, 0xF4, 0xEA),
        "note":            RGBColor(0xE8, 0xF0, 0xFE),
        "caution":         RGBColor(0xFF, 0xF3, 0xE0),
        "congratulations": RGBColor(0xF3, 0xE5, 0xF5),
        "tip":             RGBColor(0xE0, 0xF7, 0xFA),
        "team_challenge":  RGBColor(0xFF, 0xFB, 0xEB),
    }

    def _shade_paragraph(p, rgb: RGBColor):
        """Apply a background shading colour to a paragraph."""
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        pPr.append(shd)

    class _HtmlToDocx(HTMLParser):
        """Parse Quill/HTML into docx paragraphs with proper formatting."""
        def __init__(self):
            super().__init__()
            self._bold = 0
            self._italic = 0
            self._code = 0
            self._link = ""
            self._list_type = []   # stack: 'ul' | 'ol'
            self._ol_counter = []  # counter per ol level
            self._in_table = False
            self._in_th = False
            self._in_td = False
            self._table = None
            self._table_row = None
            self._table_cell = None
            self._cell_buf = []
            self._in_pre = False
            self._pre_buf = []
            self._skip_tags = {"head", "style", "script"}
            self._skip = 0
            self._cur_para = None
            self._heading_level = 0

        def _new_para(self, style="Normal"):
            self._cur_para = doc.add_paragraph(style=style)
            return self._cur_para

        def _run(self, text):
            if not text:
                return
            if self._in_pre:
                self._pre_buf.append(text)
                return
            if self._in_td or self._in_th:
                self._cell_buf.append(text)
                return
            if self._cur_para is None:
                self._cur_para = doc.add_paragraph()
            run = self._cur_para.add_run(text)
            if self._bold:
                run.bold = True
            if self._italic:
                run.italic = True
            if self._code:
                run.font.name = "Courier New"
                run.font.size = Pt(9)

        def handle_starttag(self, tag, attrs):
            attr = dict(attrs)
            if tag in self._skip_tags:
                self._skip += 1
                return
            if self._skip:
                return
            if tag in ("p", "div") and not self._in_table:
                self._cur_para = doc.add_paragraph()
            elif tag == "br":
                if self._cur_para:
                    self._cur_para.add_run("\n")
            elif tag in ("h1","h2","h3","h4","h5","h6"):
                lvl = int(tag[1])
                self._heading_level = lvl
                # Map h1→2, h2→3, h3→3 to avoid overriding section headings
                docx_lvl = min(lvl + 1, 4)
                self._cur_para = doc.add_heading("", level=docx_lvl)
            elif tag == "strong" or tag == "b":
                self._bold += 1
            elif tag == "em" or tag == "i":
                self._italic += 1
            elif tag == "code":
                self._code += 1
            elif tag == "a":
                self._link = attr.get("href", "")
            elif tag == "ul":
                self._list_type.append("ul")
            elif tag == "ol":
                self._list_type.append("ol")
                self._ol_counter.append(0)
            elif tag == "li":
                if self._list_type and self._list_type[-1] == "ol":
                    self._ol_counter[-1] += 1
                    self._cur_para = doc.add_paragraph(style="List Number")
                else:
                    self._cur_para = doc.add_paragraph(style="List Bullet")
            elif tag == "blockquote":
                self._cur_para = doc.add_paragraph(style="Quote")
            elif tag == "pre":
                self._in_pre = True
                self._pre_buf = []
            elif tag == "hr":
                # Add a horizontal line paragraph
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), 'AAAAAA')
                pBdr.append(bottom)
                pPr.append(pBdr)
                self._cur_para = None
            elif tag == "table":
                self._in_table = True
                self._table = doc.add_table(rows=0, cols=1)
                self._table.style = "Table Grid"
            elif tag == "tr":
                if self._table:
                    self._table_row = self._table.add_row()
            elif tag in ("th", "td"):
                self._cell_buf = []
                self._in_th = (tag == "th")
                self._in_td = (tag == "td")

        def handle_endtag(self, tag):
            if tag in self._skip_tags:
                self._skip = max(0, self._skip - 1)
                return
            if tag == "strong" or tag == "b":
                self._bold = max(0, self._bold - 1)
            elif tag == "em" or tag == "i":
                self._italic = max(0, self._italic - 1)
            elif tag == "code":
                self._code = max(0, self._code - 1)
            elif tag == "a":
                self._link = ""
            elif tag == "ul":
                if self._list_type:
                    self._list_type.pop()
                self._cur_para = None
            elif tag == "ol":
                if self._list_type:
                    self._list_type.pop()
                if self._ol_counter:
                    self._ol_counter.pop()
                self._cur_para = None
            elif tag == "pre":
                text = "".join(self._pre_buf).strip()
                if text:
                    p = doc.add_paragraph(text)
                    p.style = "Normal"
                    for run in p.runs:
                        run.font.name = "Courier New"
                        run.font.size = Pt(8)
                    # Light grey background
                    _shade_paragraph(p, RGBColor(0xF4, 0xF4, 0xF4))
                self._in_pre = False
                self._pre_buf = []
            elif tag in ("th", "td"):
                text = "".join(self._cell_buf).strip()
                if self._table_row:
                    col_idx = len(self._table_row.cells) - 1
                    # Expand columns if needed
                    while len(self._table.columns) <= col_idx:
                        # Add a column by adding a cell to each row
                        pass
                    try:
                        cell = self._table_row.cells[col_idx] if col_idx < len(self._table_row.cells) else self._table_row.cells[-1]
                        cell.text = text
                        if self._in_th:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
                    except Exception:
                        pass
                self._in_th = False
                self._in_td = False
                self._cell_buf = []
            elif tag == "table":
                self._in_table = False
                self._table = None
                self._table_row = None
                self._cur_para = None
            elif tag in ("h1","h2","h3","h4","h5","h6"):
                self._heading_level = 0
                self._cur_para = None
            elif tag in ("p", "li", "blockquote", "div"):
                self._cur_para = None

        def handle_data(self, data):
            data = data.replace("\u283f", "").replace("\u00a0", " ")
            if self._skip:
                return
            self._run(data)

        def handle_entityref(self, name):
            self._run(_html_mod2.unescape(f"&{name};"))

        def handle_charref(self, name):
            self._run(_html_mod2.unescape(f"&#{name};"))

    def _add_html_content(html_content: str):
        """Parse HTML content and add to docx."""
        if not html_content or not html_content.strip():
            return
        # Handle <!--markdown--> blocks — strip marker, convert to plain text
        if html_content.startswith("<!--markdown-->"):
            md = html_content[len("<!--markdown-->"):].strip()
            # Strip any inline HTML tags from markdown content
            import re as _re2
            md_clean = _re2.sub(r'<[^>]+>', '', md)
            md_clean = _html_mod2.unescape(md_clean)
            # Split on double newlines and add as paragraphs
            for chunk in md_clean.split("\n\n"):
                chunk = chunk.strip()
                if not chunk:
                    continue
                # Detect headings
                if chunk.startswith("### "):
                    doc.add_heading(chunk[4:], level=3)
                elif chunk.startswith("## "):
                    doc.add_heading(chunk[3:], level=2)
                elif chunk.startswith("# "):
                    doc.add_heading(chunk[2:], level=2)
                elif chunk.startswith("---"):
                    p = doc.add_paragraph()
                    pPr = p._p.get_or_add_pPr()
                    pBdr = OxmlElement('w:pBdr')
                    bottom = OxmlElement('w:bottom')
                    bottom.set(qn('w:val'), 'single')
                    bottom.set(qn('w:sz'), '6')
                    bottom.set(qn('w:space'), '1')
                    bottom.set(qn('w:color'), 'AAAAAA')
                    pBdr.append(bottom)
                    pPr.append(pBdr)
                elif chunk.startswith("```"):
                    code = _re2.sub(r'^```[^\n]*\n?', '', chunk).rstrip('`').strip()
                    if code:
                        p = doc.add_paragraph(code)
                        _shade_paragraph(p, RGBColor(0xF4, 0xF4, 0xF4))
                        for run in p.runs:
                            run.font.name = "Courier New"
                            run.font.size = Pt(8)
                elif _re2.match(r'^[-*] ', chunk.split('\n')[0]):
                    for line in chunk.split('\n'):
                        line = _re2.sub(r'^[-*] ', '', line).strip()
                        if line:
                            doc.add_paragraph(line, style="List Bullet")
                elif _re2.match(r'^\d+\. ', chunk.split('\n')[0]):
                    for line in chunk.split('\n'):
                        line = _re2.sub(r'^\d+\. ', '', line).strip()
                        if line:
                            doc.add_paragraph(line, style="List Number")
                elif '|' in chunk and chunk.count('|') > 2:
                    # Markdown table — render as plain text rows
                    rows = [r for r in chunk.split('\n') if r.strip() and not _re2.match(r'^\s*\|[-| ]+\|\s*$', r)]
                    if rows:
                        cols = [c.strip() for c in rows[0].split('|') if c.strip()]
                        tbl = doc.add_table(rows=len(rows), cols=len(cols))
                        tbl.style = "Table Grid"
                        for ri, row in enumerate(rows):
                            cells = [c.strip() for c in row.split('|') if c.strip()]
                            for ci, cell_text in enumerate(cells[:len(cols)]):
                                cell = tbl.rows[ri].cells[ci]
                                cell.text = cell_text
                                if ri == 0:
                                    for run in cell.paragraphs[0].runs:
                                        run.bold = True
                else:
                    # Remove markdown bold/italic markers for plain para
                    plain = _re2.sub(r'\*\*(.+?)\*\*', r'\1', chunk)
                    plain = _re2.sub(r'\*(.+?)\*', r'\1', plain)
                    plain = _re2.sub(r'__(.+?)__', r'\1', plain)
                    plain = _re2.sub(r'`(.+?)`', r'\1', plain)
                    if plain.strip():
                        doc.add_paragraph(plain.strip())
            return
        # Regular HTML — use parser
        parser = _HtmlToDocx()
        parser.feed(html_content)

    def _add_callout(callout_type: str, content: str, caption: str = ""):
        label = caption if caption and caption.strip() else _CALLOUT_LABEL.get(callout_type, callout_type or "Note")
        color = _CALLOUT_COLOR.get(callout_type, RGBColor(0xE8, 0xF0, 0xFE))
        # Label paragraph
        p_label = doc.add_paragraph()
        run = p_label.add_run(f"  {label}")
        run.bold = True
        _shade_paragraph(p_label, color)
        # Content paragraph(s)
        text = _to_plain(content) if content else ""
        for chunk in text.split("\n\n"):
            chunk = chunk.strip()
            if chunk:
                p = doc.add_paragraph(f"  {chunk}")
                _shade_paragraph(p, color)

    def _add_screenshot(path_str: str):
        img_path = _resolve_screenshot_path(path_str)
        if img_path.exists():
            try:
                doc.add_picture(str(img_path), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

    def _add_blocks(blocks):
        for blk in blocks:
            if blk.type == "text" and blk.content:
                _add_html_content(blk.content)
            elif blk.type == "screenshot" and blk.path:
                _add_screenshot(blk.path)
            elif blk.type == "callout":
                _add_callout(blk.callout_type or "note", blk.content or "", blk.caption or "")
            elif blk.type == "divider":
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '6')
                bottom.set(qn('w:space'), '1')
                bottom.set(qn('w:color'), 'AAAAAA')
                pBdr.append(bottom)
                pPr.append(pBdr)

    doc = Document()
    m = guide.metadata

    # ── Document styles tweaks ───────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    doc.add_heading(m.title.strip(), 0)
    meta_para = doc.add_paragraph(
        f"Version {m.version}  |  {m.author or 'Unknown'}  |  {m.date}  |  "
        f"{m.difficulty.capitalize()}  |  {m.lab_duration_minutes} min"
    )
    try:
        meta_para.style = doc.styles["Caption"]
    except Exception:
        pass

    if m.prerequisites:
        doc.add_heading("Prerequisites", 2)
        for prereq in m.prerequisites:
            doc.add_paragraph(prereq, style="List Bullet")

    if guide.introduction:
        doc.add_heading("Introduction", 1)
        intro_text = _to_plain(guide.introduction)
        for chunk in intro_text.split("\n\n"):
            chunk = chunk.strip()
            if chunk:
                doc.add_paragraph(chunk)

    if guide.learning_objectives:
        doc.add_heading("Learning Objectives", 1)
        for obj in guide.learning_objectives:
            doc.add_paragraph(obj.text, style="List Bullet")

    for sec in guide.sections:
        doc.add_heading(sec.title, 1)
        if sec.overview:
            doc.add_paragraph(sec.overview)

        sec_blocks = getattr(sec, "blocks", []) or []
        if sec_blocks:
            _add_blocks(sec_blocks)
        else:
            for ss in getattr(sec, "screenshots", []):
                _add_screenshot(ss.path)

        for step in sec.steps:
            doc.add_heading(f"Step {step.order}: {step.title}", 2)
            if step.instruction and step.instruction.strip():
                _add_html_content(step.instruction)

            for cb in step.code_blocks:
                p = doc.add_paragraph(cb)
                _shade_paragraph(p, RGBColor(0xF4, 0xF4, 0xF4))
                for run in p.runs:
                    run.font.name = "Courier New"
                    run.font.size = Pt(8)

            step_blocks = getattr(step, "blocks", []) or []
            if step_blocks:
                _add_blocks(step_blocks)
            else:
                for ss in step.screenshots:
                    _add_screenshot(ss.path)

            if getattr(step, "expected_result", ""):
                _add_callout("expected_result", step.expected_result)
            if getattr(step, "notes", ""):
                _add_callout("note", step.notes)

    if guide.conclusion:
        doc.add_heading("Conclusion", 1)
        conc_text = _to_plain(guide.conclusion)
        for chunk in conc_text.split("\n\n"):
            chunk = chunk.strip()
            if chunk:
                doc.add_paragraph(chunk)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ─────────────────────────────────────────────────────────────
# MkDocs site
# ─────────────────────────────────────────────────────────────

_MKDOCS_YML_TEMPLATE = """\
---
dev_addr: "0.0.0.0:8000"
site_name: "{title}"
site_description: "{title}"
site_author: "Cisco"
copyright: Copyright &copy; 2024 Cisco
theme:
  name: material
  features:
    - navigation.indexes
    - navigation.instant
    - navigation.top
    - navigation.footer
    - navigation.expand
    - search.suggest
    - content.code.copy
  custom_dir: docs/overrides
  palette:
    - media: "(prefers-color-scheme: light)"
      scheme: default
      primary: custom
      toggle:
        icon: material/brightness-7
        name: Switch to dark mode
    - media: "(prefers-color-scheme: dark)"
      scheme: slate
      primary: custom
      toggle:
        icon: material/brightness-4
        name: Switch to light mode
  logo: template_assets/cisco_logo.png
  favicon: template_assets/cisco_logo.png
extra:
  generator: false
extra_css:
  - stylesheets/extra.css
extra_javascript:
  - javascripts/links.js
plugins:
  - search
  - glightbox:
      touchNavigation: true
      loop: false
      effect: fade
      slide_effect: slide
      width: 100%
      height: auto
      zoomable: true
      draggable: false
      auto_caption: true
      caption_position: top
markdown_extensions:
  - abbr
  - admonition
  - attr_list
  - def_list
  - footnotes
  - meta
  - md_in_html
  - tables
  - toc:
      toc_depth: 2
  - pymdownx.details
  - pymdownx.highlight
  - pymdownx.superfences:
      custom_fences:
        - name: mermaid
          class: mermaid
          format: !!python/name:pymdownx.superfences.fence_code_format
nav:
  - Home: index.md
{nav_sections}
"""

_EXTRA_CSS = """\
.md-typeset__table {
    min-width: 100%;
}

.md-typeset table:not([class]) {
    display: table;
}

.md-grid {
    max-width: 1440px;
}

:root {
    --md-primary-fg-color: rgb(11, 23, 44);
    --md-accent-fg-color: #00bdeb;
}

[data-md-color-scheme="default"] {
    --md-typeset-a-color: #3a7fff;
    --md-accent-fg-color: #00bdeb;
}

[data-md-color-scheme="slate"] {
    --md-typeset-a-color: #3a7fff;
    --md-accent-fg-color: #00bdeb;
}

/* ── Team Challenge custom admonition ── */
:root {
    --md-admonition-icon--team-challenge: url('data:image/svg+xml;charset=utf-8,<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 24 24"><path d="M11.049 2.927c.3-.921 1.603-.921 1.902 0l1.519 4.674a1 1 0 0 0 .95.69h4.915c.969 0 1.371 1.24.588 1.81l-3.976 2.888a1 1 0 0 0-.363 1.118l1.518 4.674c.3.922-.755 1.688-1.538 1.118l-3.976-2.888a1 1 0 0 0-1.176 0l-3.976 2.888c-.783.57-1.838-.197-1.538-1.118l1.518-4.674a1 1 0 0 0-.363-1.118l-3.976-2.888c-.784-.57-.38-1.81.588-1.81h4.914a1 1 0 0 0 .951-.69l1.519-4.674z"/></svg>');
}
.md-typeset .admonition.team-challenge,
.md-typeset details.team-challenge {
    border-color: #f59e0b;
    background-color: #fffbeb;
}
[data-md-color-scheme="slate"] .md-typeset .admonition.team-challenge,
[data-md-color-scheme="slate"] .md-typeset details.team-challenge {
    background-color: #2a1f00;
}
.md-typeset .team-challenge > .admonition-title,
.md-typeset .team-challenge > summary {
    background-color: rgba(245, 158, 11, 0.15);
    border-color: #f59e0b;
    color: #b45309;
}
[data-md-color-scheme="slate"] .md-typeset .team-challenge > .admonition-title,
[data-md-color-scheme="slate"] .md-typeset .team-challenge > summary {
    color: #fcd34d;
}
.md-typeset .team-challenge > .admonition-title::before,
.md-typeset .team-challenge > summary::before {
    background-color: #f59e0b;
    -webkit-mask-image: var(--md-admonition-icon--team-challenge);
    mask-image: var(--md-admonition-icon--team-challenge);
}

/* ── Code block sizing — keep long prompts readable and contained ── */
.md-typeset pre {
    font-size: .65rem;
    line-height: 1.6;
    white-space: pre-wrap;
    word-break: break-word;
    max-height: none !important;
    overflow-y: visible !important;
    overflow-x: visible !important;
}
.md-typeset code {
    font-size: .68rem;
}

/* ── Callout admonition emoji animations ── */
@keyframes callout-pulse  { 0%,100%{opacity:1;transform:scale(1)}    50%{opacity:.4;transform:scale(.85)} }
@keyframes callout-shake  { 0%,100%{transform:rotate(0)}  20%{transform:rotate(-15deg)}  40%{transform:rotate(15deg)}  60%{transform:rotate(-10deg)}  80%{transform:rotate(10deg)} }
@keyframes callout-bounce { 0%,100%{transform:translateY(0)}  40%{transform:translateY(-5px)}  60%{transform:translateY(-2px)} }
@keyframes callout-spin   { 0%{transform:rotate(0deg)} 100%{transform:rotate(360deg)} }
@keyframes callout-pop    { 0%,100%{transform:scale(1)}  50%{transform:scale(1.35)} }
@keyframes callout-wobble { 0%,100%{transform:rotate(0) scale(1)}  25%{transform:rotate(-8deg) scale(1.1)}  75%{transform:rotate(8deg) scale(1.1)} }

.md-typeset .admonition.note > .admonition-title::before,
.md-typeset .admonition.info > .admonition-title::before
  { animation: callout-bounce 2.2s ease-in-out infinite; }

.md-typeset .admonition.warning > .admonition-title::before,
.md-typeset .admonition.caution > .admonition-title::before
  { animation: callout-shake 2.0s ease-in-out infinite; }

.md-typeset .admonition.success > .admonition-title::before,
.md-typeset .admonition.check > .admonition-title::before
  { animation: callout-pulse 1.8s ease-in-out infinite; }

.md-typeset .admonition.abstract > .admonition-title::before
  { animation: callout-pop 1.5s ease-in-out infinite; }

.md-typeset .admonition.tip > .admonition-title::before,
.md-typeset .admonition.hint > .admonition-title::before
  { animation: callout-spin 3.0s linear infinite; }

.md-typeset .admonition.team-challenge > .admonition-title::before
  { animation: callout-wobble 1.8s ease-in-out infinite; }

/* ── Hide right-hand TOC sidebar, expand content area ── */
.md-sidebar--secondary,
.md-sidebar--secondary * {
    display: none !important;
    width: 0 !important;
    min-width: 0 !important;
}
.md-content {
    max-width: none !important;
    margin-right: 0 !important;
}
@media screen and (min-width: 60em) {
    .md-content {
        margin-right: 0 !important;
    }
}
@media screen and (min-width: 76.25em) {
    .md-content {
        margin-right: 0 !important;
    }
}
"""

_HOME_HTML = """\
{% extends "main.html" %}
{% block tabs %}
{{ super() }}
<style>
    .md-container { height: 100%; background: var(--md-primary-fg-color) }
    .md-main { flex-grow: 0; height: 0px; background: var(--md-primary-fg-color) }
    .md-main__inner { display: flex; height: 100%; }
    .tx-container { padding-top: .0rem; background: var(--md-primary-fg-color) }
    .tx-hero { margin: 32px 2.8rem; color: var(--md-primary-bg-color); justify-content: center; }
    .tx-hero h1 { margin-bottom: 1rem; color: currentColor; font-weight: 700 }
    :root { .tx-hero__content { margin: 0; } }
    .tx-hero__content { padding-bottom: 1rem; margin: 0 auto; }
    .tx-hero__logo { scale: 1; width: 100%; padding-bottom: 50px; }
    .tx-hero__image { width: 850px; height: 1050px; order: 1; padding-right: 2.5rem; }
    .tx-hero .md-button { margin-top: .5rem; margin-right: .5rem; color: var(--md-primary-bg-color) }
    .tx-hero .md-button--primary { background-color: var(--md-primary-bg-color); color: rgb(0,0,0); border-color: var(--md-primary-bg-color) }
    .tx-hero .md-button:focus, .tx-hero .md-button:hover { background-color: var(--md-accent-fg-color); color: var(--md-default-bg-color); border-color: var(--md-accent-fg-color) }
    @media screen and (min-width:60em) {
        .md-sidebar--secondary { display: none }
        .tx-hero { display: flex; align-items: center; justify-content: center; }
        .tx-hero__content { max-width: 22rem; margin-top: 3.5rem; margin-bottom: 3.5rem; margin-left: 1.0rem; margin-right: 4.0rem; align-items: center; }
    }
    @media screen and (min-width:76.25em) {
        .md-sidebar--primary { display: none }
    }
</style>
<section class="tx-container">
    <div class="md-grid md-typeset">
        <div class="tx-hero">
            <div class="tx-hero__image">
                <img src="template_assets/CLAMER2025_Static_Midnight_Generic.png" draggable="false">
            </div>
            <div class="tx-hero__content">
                <div class="tx-hero__logo">
                    <img src="template_assets/TE_white_Logo_300dpi.png" draggable="false">
                </div>
                <h1>
                    <script>
                        var title = document.querySelector('meta[name="description"]').content
                        const lab = title.split(" ")[0]
                        document.write(lab)
                    </script>
                </h1>
                <p>
                    <script>
                        const descr = title.substring(title.indexOf(" ") + 1)
                        document.write(descr)
                    </script>
                </p>
                <a href="{{ nav.items[1].url }}" class="md-button md-button--primary">
                    Get started
                </a>
            </div>
        </div>
    </div>
</section>
{% endblock %}
{% block content %} {% endblock %}
{% block footer %} {% endblock %}
"""

_GITHUB_ACTIONS_WORKFLOW = """\
name: Deploy Lab Guide to GitHub Pages

on:
  push:
    branches:
      - main

permissions:
  contents: write

jobs:
  deploy:
    runs-on: ubuntu-latest
    steps:
      - uses: actions/checkout@v4

      - name: Set up Python
        uses: actions/setup-python@v5
        with:
          python-version: '3.11'

      - name: Install MkDocs dependencies
        run: pip install mkdocs-material mkdocs-glightbox

      - name: Deploy to GitHub Pages
        run: mkdocs gh-deploy --force
"""

# Reference assets bundled from ciscodocs/ltrxar-3783-tech-elevate-fy26
_REFERENCE_ASSETS_REPO = "https://github.com/ciscodocs/ltrxar-3783-tech-elevate-fy26.git"

_CALLOUT_ADMONITION = {
    "expected_result": ('success',  "Expected Result"),
    "note":            ('info',     "Note"),
    "caution":         ('warning',  "Caution"),
    "congratulations": ('abstract', "Congratulations"),
    "tip":             ('tip',      "Tip"),
    "team_challenge":  ('team-challenge', "🏆 Team Challenge"),
}


def _render_block_mkdocs(blk) -> list[str]:
    """Return MkDocs markdown lines for a single ContentBlock."""
    lines: list[str] = []
    if blk.type == "divider":
        lines += ["---", ""]
    if blk.type == "text" and blk.content:
        content = blk.content
        # If content is pre-rendered Markdown (marked with <!--markdown-->), pass through as-is
        if content.startswith("<!--markdown-->"):
            lines += [content[len("<!--markdown-->"):].lstrip("\n"), ""]
        else:
            lines += [_html_to_md(content), ""]
    elif blk.type == "screenshot" and blk.path:
        fname = Path(blk.path).name
        lines += [f"![](../screenshots/{fname})", ""]
    elif blk.type == "callout":
        kind, title = _CALLOUT_ADMONITION.get(blk.callout_type, ("note", blk.callout_type))
        if getattr(blk, "caption", "") and blk.caption.strip():
            title = blk.caption
        body = _html_to_md(blk.content or "")
        # Every line must be indented 4 spaces for MkDocs admonition to contain it
        indented = "\n".join("    " + line if line.strip() else "" for line in body.splitlines())
        lines += [f'!!! {kind} "{title}"', indented, ""]
    return lines


def _render_index_md(guide: LabGuide) -> str:
    """Render the home page index.md — intro only, with hero template frontmatter."""
    m = guide.metadata
    lines = [
        "---",
        f'description: "{m.title}"',
        "template: home.html",
        "---",
        "",
        f"# {m.title}",
        "",
    ]
    if m.subtitle:
        lines += [m.subtitle, ""]
    if m.author or m.version:
        bits = []
        if m.author:
            bits.append(f"**Author:** {m.author}")
        if m.version:
            bits.append(f"**Version:** {m.version}")
        lines += ["  |  ".join(bits), ""]
    if guide.learning_objectives:
        lines += ["## Learning Objectives", ""]
        for obj in guide.learning_objectives:
            lines += [f"- {obj.text}"]
        lines += [""]
    return "\n".join(lines)


def export_mkdocs(guide: LabGuide, output_dir: Path) -> Path:
    """
    Generate a full MkDocs docs/ directory structure.
    output_dir/
      mkdocs.yml
      docs/
        index.md
        <section-slug>/
          index.md
    """
    docs_dir = output_dir / "docs"
    docs_dir.mkdir(parents=True, exist_ok=True)

    # ── Clean stale section dirs from previous exports ──────────
    import shutil as _shutil_clean
    import re as _re_slug
    def slugify(text: str) -> str:
        return _re_slug.sub(r"[^\w]+", "-", text.lower()).strip("-")

    # Compute current slugs so we only keep them + reserved dirs
    current_slugs = {slugify(sec.title) for sec in guide.sections}
    reserved_dirs = {"screenshots", "stylesheets", "overrides", "template_assets", "introduction", "conclusion", "javascripts"}
    for child in list(docs_dir.iterdir()):
        if child.is_dir() and child.name not in reserved_dirs and child.name not in current_slugs:
            _shutil_clean.rmtree(child)

    section_slugs = [(slugify(sec.title), sec.title) for sec in guide.sections]

    # index.md — hero landing page (no introduction content here)
    (docs_dir / "index.md").write_text(_render_index_md(guide))

    # Copy screenshots into docs/screenshots/ so MkDocs can serve them
    _copy_screenshots(guide, docs_dir)

    # ── Introduction page ────────────────────────────────────────
    nav_intro = ""
    if guide.introduction:
        intro_dir = docs_dir / "introduction"
        intro_dir.mkdir(exist_ok=True)
        intro_lines = ["# Introduction", "", guide.introduction, ""]
        if guide.learning_objectives:
            intro_lines += ["## Learning Objectives", ""]
            for obj in guide.learning_objectives:
                intro_lines += [f"- {obj.text}"]
            intro_lines += [""]
        (intro_dir / "index.md").write_text("\n".join(intro_lines))
        nav_intro = "  - Introduction: introduction/index.md\n"

    # Per-section pages
    nav_sections = ""
    for idx, sec in enumerate(guide.sections):
        slug, _ = section_slugs[idx]
        sec_dir = docs_dir / slug
        sec_dir.mkdir(exist_ok=True)

        lines = [f"# {sec.title}", ""]
        if sec.overview:
            import re as _re
            fixed_overview = _re.sub(
                r'!\[([^\]]*)\]\(screenshots/([^)]+)\)',
                r'![\1](../screenshots/\2)',
                sec.overview
            )
            lines += [fixed_overview, ""]

        # Section blocks (new system)
        sec_blocks = getattr(sec, "blocks", []) or []
        if sec_blocks:
            for blk in sec_blocks:
                lines += _render_block_mkdocs(blk)
        else:
            for ss in getattr(sec, "screenshots", []):
                fname = Path(ss.path).name
                lines += [f"![](../screenshots/{fname})", ""]

        for step in sec.steps:
            # Skip "Step N:" prefix for appendix sections
            is_appendix = "appendix" in sec.title.lower()
            step_heading = step.title if is_appendix else f"Step {step.order}: {step.title}"
            lines += [f"## {step_heading}", ""]
            _instr = step.instruction or ""
            if _instr.strip().startswith("<!--markdown-->"):
                lines += [_instr.strip()[len("<!--markdown-->"):].lstrip("\n"), ""]
            else:
                lines += [_html_to_md(_instr), ""]
            for cb in step.code_blocks:
                lines += ["```", cb, "```", ""]

            # Step blocks (new system)
            step_blocks = getattr(step, "blocks", []) or []
            if step_blocks:
                for blk in step_blocks:
                    lines += _render_block_mkdocs(blk)
            else:
                for ss in step.screenshots:
                    fname = Path(ss.path).name
                    lines += [f"![](../screenshots/{fname})", ""]
            # Legacy dedicated fields
            if getattr(step, "expected_result", ""):
                lines += ['!!! success "Expected Result"', f"    {step.expected_result}", ""]
            if getattr(step, "notes", ""):
                lines += ['!!! note "Note"', f"    {step.notes}", ""]

        (sec_dir / "index.md").write_text("\n".join(lines))
        # Quote the label if it contains a colon (would break YAML otherwise)
        label = f'"{sec.title}"' if ":" in sec.title else sec.title
        nav_sections += f"  - {label}: {slug}/index.md\n"

    # ── Conclusion page ──────────────────────────────────────────
    nav_conclusion = ""
    if guide.conclusion:
        conc_dir = docs_dir / "conclusion"
        conc_dir.mkdir(exist_ok=True)
        conc_lines = ["# Conclusion", "", guide.conclusion, ""]
        (conc_dir / "index.md").write_text("\n".join(conc_lines))
        nav_conclusion = "  - Conclusion: conclusion/index.md\n"

    # mkdocs.yml
    yml = _MKDOCS_YML_TEMPLATE.format(
        title=guide.metadata.title,
        nav_sections=nav_intro + nav_sections + nav_conclusion,
    )
    (output_dir / "mkdocs.yml").write_text(yml)

    # ── Cisco brand assets ──────────────────────────────────────
    _write_brand_assets(docs_dir)

    # ── GitHub Actions workflow ─────────────────────────────────
    workflow_dir = output_dir / ".github" / "workflows"
    workflow_dir.mkdir(parents=True, exist_ok=True)
    (workflow_dir / "deploy.yml").write_text(_GITHUB_ACTIONS_WORKFLOW)

    return output_dir


def _copy_screenshots(guide: LabGuide, docs_dir: Path) -> None:
    """Copy all referenced screenshots into docs_dir/screenshots/."""
    import shutil as _shutil
    ss_dst = docs_dir / "screenshots"
    ss_dst.mkdir(exist_ok=True)
    ss_src = Path(__file__).parent.parent / "data" / "screenshots"
    for sec in guide.sections:
        # Block-level screenshots
        for blk in getattr(sec, "blocks", []):
            if blk.type == "screenshot" and blk.path:
                # section blocks store path as "screenshots/filename"
                src = ss_src / Path(blk.path).name
                if src.exists():
                    _shutil.copy2(src, ss_dst / src.name)
        # Legacy section-level screenshots
        for ss in getattr(sec, "screenshots", []):
            src = ss_src / Path(ss.path).name
            if src.exists():
                _shutil.copy2(src, ss_dst / src.name)
        # Step-level screenshots and blocks
        for step in sec.steps:
            for blk in getattr(step, "blocks", []):
                if blk.type == "screenshot" and blk.path:
                    # step blocks store path as bare filename
                    src = ss_src / Path(blk.path).name
                    if src.exists():
                        _shutil.copy2(src, ss_dst / src.name)
            for ss in step.screenshots:
                src = ss_src / Path(ss.path).name
                if src.exists():
                    _shutil.copy2(src, ss_dst / src.name)


def _write_brand_assets(docs_dir: Path) -> None:
    """Write Cisco brand assets (CSS, overrides, images) into docs_dir."""
    import shutil, base64

    # Stylesheets
    (docs_dir / "stylesheets").mkdir(exist_ok=True)
    (docs_dir / "stylesheets" / "extra.css").write_text(_EXTRA_CSS)

    # Javascripts
    (docs_dir / "javascripts").mkdir(exist_ok=True)
    (docs_dir / "javascripts" / "links.js").write_text(
        "document.addEventListener('DOMContentLoaded', function () {\n"
        "  document.querySelectorAll('a[href]').forEach(function (a) {\n"
        "    if (!a.getAttribute('target')) {\n"
        "      a.setAttribute('target', '_blank');\n"
        "      a.setAttribute('rel', 'noopener noreferrer');\n"
        "    }\n"
        "  });\n"
        "});\n"
    )

    # Overrides
    (docs_dir / "overrides").mkdir(exist_ok=True)
    (docs_dir / "overrides" / "home.html").write_text(_HOME_HTML)

    # Template assets — try multiple source locations in priority order
    _ASSET_SEARCH_PATHS = [
        Path("/tmp/ltrxar-3783/docs/template_assets"),
        Path("/Users/maokuma/Documents/CODE/ltrxar-3783-clamer25/docs/template_assets"),
        Path("/Users/maokuma/sw_projects/lab_guide_automator/data/exports/84ff424d-ec65-4579-8be3-b9b7101bf6ea-mkdocs/docs/template_assets"),
    ]
    asset_dst = docs_dir / "template_assets"
    asset_dst.mkdir(exist_ok=True)
    needed = [
        "cisco_logo.png",
        "TE_white_Logo_300dpi.png",
        "CLAMER2025_Static_Midnight_Generic.png",
    ]
    for fname in needed:
        dst = asset_dst / fname
        if dst.exists():
            continue
        for search_path in _ASSET_SEARCH_PATHS:
            src = search_path / fname
            if src.exists():
                shutil.copy2(src, dst)
                break


def push_mkdocs_to_git(output_dir: Path, repo_url: str, branch: str = "main") -> str:
    """
    Build/update a git repo in output_dir and push to repo_url.

    Injects the mokuma56 PAT into the HTTPS URL so the push bypasses
    the macOS keychain (which returns the EMU account instead).
    """
    import re as _re

    # Inject credentials into HTTPS URL: https://user:token@github.com/...
    _TOKEN = "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"
    _USER  = "mokuma56"
    if repo_url.startswith("https://github.com/"):
        auth_url = repo_url.replace("https://", f"https://{_USER}:{_TOKEN}@")
    else:
        auth_url = repo_url   # SSH or custom — use as-is

    def _git(*args, check=True):
        return subprocess.run(
            ["git", *args],
            cwd=output_dir,
            check=check,
            capture_output=True,
            text=True,
        )

    git_dir = output_dir / ".git"
    # Check that .git belongs to output_dir itself, not a parent repo
    is_fresh = not git_dir.exists()
    if not is_fresh:
        # Verify this git repo is rooted at output_dir, not a parent
        result = subprocess.run(
            ["git", "rev-parse", "--show-toplevel"],
            cwd=output_dir, capture_output=True, text=True
        )
        if result.stdout.strip() != str(output_dir.resolve()):
            is_fresh = True  # parent repo — treat as fresh, init our own

    # ── Init if needed ──────────────────────────────────────────
    if is_fresh:
        _git("init")
        _git("remote", "add", "origin", auth_url)
        # Fetch so --force-with-lease has ref tracking info
        _git("fetch", "origin", check=False)
    else:
        result = _git("remote", "get-url", "origin", check=False)
        if result.returncode != 0:
            _git("remote", "add", "origin", auth_url)
        else:
            _git("remote", "set-url", "origin", auth_url)
        # Always fetch so --force-with-lease has current ref tracking info
        _git("fetch", "origin", check=False)

    # ── Ensure a git identity exists (local fallback) ───────────
    for key, val in [("user.email", "lab-guide@localhost"), ("user.name", "Lab Guide Automator")]:
        r = _git("config", key, check=False)
        if not r.stdout.strip():
            _git("config", key, val)

    # ── Stage + commit (always, so GitHub receives a push event for Actions) ──
    _git("add", "-A")
    status = _git("status", "--porcelain")
    if status.stdout.strip():
        _git("commit", "-m", f"Update lab guide — {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    else:
        # Nothing changed in content but we still need to trigger GitHub Actions,
        # so create an empty commit to generate a new push event.
        _git("commit", "--allow-empty", "-m", f"Trigger deploy — {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # ── Push ────────────────────────────────────────────────────
    _git("push", "-u", "--force", "origin", branch)
    return f"Pushed to {repo_url} ({branch})"


# ─────────────────────────────────────────────────────────────
# Narrated video (TTS + ffmpeg)
# ─────────────────────────────────────────────────────────────

def _build_narration_script(guide: LabGuide) -> str:
    """Build a narration script from the guide (one sentence per step)."""
    lines = [f"Welcome to the lab: {guide.metadata.title}. "]
    if guide.introduction:
        lines.append(guide.introduction.split("\n")[0])
    for sec in guide.sections:
        lines.append(f"Section: {sec.title}. ")
        for step in sec.steps:
            lines.append(f"Step {step.order}: {step.title}. {step.instruction.split('.')[0]}. ")
    if guide.conclusion:
        lines.append(guide.conclusion.split("\n")[0])
    return " ".join(lines)


async def generate_narration_audio(
    settings,
    guide: LabGuide,
    output_path: Path,
) -> Path:
    """Generate narration audio using the configured TTS provider."""
    script = _build_narration_script(guide)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if settings.tts_provider == "openai":
        from openai import AsyncOpenAI
        client = AsyncOpenAI(api_key=settings.openai_api_key)
        response = await client.audio.speech.create(
            model="tts-1-hd",
            voice=settings.tts_voice,
            input=script,
        )
        output_path.write_bytes(response.content)

    elif settings.tts_provider == "edge":
        # edge-tts (free, no API key needed)
        import edge_tts  # type: ignore
        communicate = edge_tts.Communicate(script, voice="en-US-AriaNeural")
        await communicate.save(str(output_path))

    else:
        # macOS system TTS fallback
        mp3_path = output_path.with_suffix(".aiff")
        subprocess.run(["say", "-o", str(mp3_path), script], check=True)
        subprocess.run(
            [_ffmpeg_bin(), "-y", "-i", str(mp3_path), str(output_path)],
            check=True, capture_output=True,
        )
        mp3_path.unlink(missing_ok=True)

    return output_path


async def export_narrated_video(
    settings,
    guide: LabGuide,
    recording_path: Path,
    output_path: Path,
    session_dir: Path,
) -> Path:
    """
    Combine the original recording with AI-generated narration audio.
    If recording has audio, it's mixed; otherwise narration is the only track.
    """
    narration_path = session_dir / "narration.mp3"
    await generate_narration_audio(settings, guide, narration_path)

    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Mix original audio (if any) with narration, or just replace
    subprocess.run([
        _ffmpeg_bin(), "-y",
        "-i", str(recording_path),
        "-i", str(narration_path),
        "-filter_complex",
        "[0:a][1:a]amix=inputs=2:duration=first:dropout_transition=2[aout]",
        "-map", "0:v",
        "-map", "[aout]",
        "-c:v", "copy",
        "-c:a", "aac",
        "-shortest",
        str(output_path),
    ], check=True, capture_output=True)

    return output_path
