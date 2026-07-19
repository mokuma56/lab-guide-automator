"""
Generate the Lab Guide Automator — User & Installation Guide as a Word .docx
Run: uv run python generate_docs.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from pathlib import Path


# ── Colour palette ────────────────────────────────────────────────────────────
CISCO_BLUE   = RGBColor(0x00, 0x50, 0x73)   # #005073
CISCO_CYAN   = RGBColor(0x00, 0xBC, 0xEB)   # #00bceb
DARK         = RGBColor(0x1A, 0x1A, 0x1A)
MID          = RGBColor(0x44, 0x44, 0x44)
LIGHT_BG     = RGBColor(0xF4, 0xF8, 0xFB)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
GREEN        = RGBColor(0x2E, 0x86, 0x48)
ORANGE       = RGBColor(0xE6, 0x7E, 0x22)


def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    """Add a thin Cisco-cyan horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '00BCEB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = WHITE
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = CISCO_BLUE
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x7A, 0x8C)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    return p


def body(doc, text, bold_parts=None):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size  = Pt(10.5)
        run.font.color.rgb = DARK
    return p


def note_box(doc, text, label="💡 Note", color="1F7A8C"):
    """Shaded note / tip box."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'E8F7FC')
    cell.width = Inches(6)
    p = cell.paragraphs[0]
    run = p.add_run(f"{label}  ")
    run.bold = True
    run.font.color.rgb = RGBColor(int(color[:2],16), int(color[2:4],16), int(color[4:],16))
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    run2.font.color.rgb = DARK
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    doc.add_paragraph()  # spacer
    return table


def warning_box(doc, text):
    note_box(doc, text, label="⚠️  Important", color="E67E22")


def step_table(doc, steps):
    """Numbered step table."""
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(0.45)
    table.columns[1].width = Inches(5.7)
    for i, (action, detail) in enumerate(steps, 1):
        row = table.add_row()
        # number cell
        num_cell = row.cells[0]
        set_cell_bg(num_cell, '005073')
        np = num_cell.paragraphs[0]
        nr = np.add_run(str(i))
        nr.bold = True
        nr.font.color.rgb = WHITE
        nr.font.size = Pt(11)
        np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # content cell
        c_cell = row.cells[1]
        cp = c_cell.paragraphs[0]
        ar = cp.add_run(action)
        ar.bold = True
        ar.font.size = Pt(10)
        ar.font.color.rgb = CISCO_BLUE
        if detail:
            cp.add_run(f"\n{detail}").font.size = Pt(10)
    doc.add_paragraph()
    return table


def code_block(doc, code_text):
    """Monospace shaded code block."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, 'F0F4F8')
    p = cell.paragraphs[0]
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    doc.add_paragraph()


def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.25 + level * 0.2)
    p.paragraph_format.space_after  = Pt(3)
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK
    return p


def bold_bullet(doc, label, detail):
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = CISCO_BLUE
    r2 = p.add_run(detail)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(3)
    return p


def feature_table(doc, rows_data, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '005073')
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
    for row_data in rows_data:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.color.rgb = DARK
    doc.add_paragraph()
    return table


# ══════════════════════════════════════════════════════════════════════════════
# Build Document
# ══════════════════════════════════════════════════════════════════════════════

doc = Document()

# Page margins
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# Default body font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)
style.font.color.rgb = DARK

# Heading styles
for lvl, sz, bold in [(1,22,True),(2,16,True),(3,13,True)]:
    s = doc.styles[f'Heading {lvl}']
    s.font.name  = 'Calibri'
    s.font.size  = Pt(sz)
    s.font.bold  = bold


# ── COVER PAGE ────────────────────────────────────────────────────────────────

# Cisco blue banner
cover_table = doc.add_table(rows=1, cols=1)
cover_table.style = 'Table Grid'
cover_cell = cover_table.cell(0,0)
set_cell_bg(cover_cell, '005073')
cp = cover_cell.paragraphs[0]
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cp.add_run('\n\nLab Guide Automator\n')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = WHITE
r2 = cp.add_run('User Guide & Installation Manual\n\n')
r2.font.size = Pt(16)
r2.font.color.rgb = CISCO_CYAN
r3 = cp.add_run('Version 1.0  ·  July 2026\n\n')
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0xCC,0xE8,0xF5)

doc.add_paragraph()

# Subtitle block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'An AI-assisted platform for creating, editing, and publishing\n'
    'professional lab guides with screenshots, structured steps,\n'
    'and one-click GitHub Pages publishing.'
)
r.font.size = Pt(12)
r.font.color.rgb = MID
r.italic = True

doc.add_paragraph()
add_horizontal_rule(doc)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Cisco Systems, Inc.  ·  Internal Use')
r.font.size = Pt(10)
r.font.color.rgb = MID

doc.add_page_break()


# ── TABLE OF CONTENTS (manual) ────────────────────────────────────────────────

h2(doc, 'Table of Contents')
add_horizontal_rule(doc)
toc = [
    ('1', 'Overview',                                     '3'),
    ('2', 'System Requirements',                          '4'),
    ('3', 'Installation',                                 '5'),
    ('  3.1', 'Prerequisites',                            '5'),
    ('  3.2', 'Clone the Repository',                     '5'),
    ('  3.3', 'Install Dependencies',                     '6'),
    ('  3.4', 'Environment Configuration',                '6'),
    ('  3.5', 'First Launch',                             '7'),
    ('4', 'Dashboard Tour',                               '8'),
    ('  4.1', 'Guide Library (Left Panel)',               '8'),
    ('  4.2', 'Content Tab',                              '8'),
    ('  4.3', 'Preview Tab',                              '9'),
    ('  4.4', 'Export Tab',                               '9'),
    ('  4.5', 'AI Tools Tab',                             '9'),
    ('5', 'Creating Your First Guide',                    '10'),
    ('6', 'Editing Guides',                               '11'),
    ('  6.1', 'Introduction & Conclusion',                '11'),
    ('  6.2', 'Sections',                                 '11'),
    ('  6.3', 'Steps',                                    '12'),
    ('  6.4', 'Content Blocks',                           '13'),
    ('7', 'Screenshots',                                  '14'),
    ('  7.1', 'Screenshot Repository',                    '14'),
    ('  7.2', 'Attaching Screenshots to Steps',           '14'),
    ('  7.3', 'AI Captions',                              '15'),
    ('8', 'Ingesting Existing Documents',                 '15'),
    ('9', 'Exporting & Publishing',                       '16'),
    ('  9.1', 'Markdown Export',                          '16'),
    ('  9.2', 'HTML Export',                              '16'),
    ('  9.3', 'MkDocs / GitHub Pages Publish',            '16'),
    ('10', 'AI Features',                                 '17'),
    ('11', 'Keyboard & Interaction Reference',            '18'),
    ('12', 'Troubleshooting',                             '19'),
    ('13', 'API Reference',                               '20'),
]

toc_table = doc.add_table(rows=0, cols=3)
toc_table.style = 'Table Grid'
toc_table.columns[0].width = Inches(0.5)
toc_table.columns[1].width = Inches(5.3)
toc_table.columns[2].width = Inches(0.4)
for num, title, page in toc:
    row = toc_table.add_row()
    for i, val in enumerate([num, title, page]):
        p = row.cells[i].paragraphs[0]
        run = p.add_run(val)
        run.font.size = Pt(10)
        is_main = not num.startswith(' ')
        run.bold = is_main and (i < 2)
        run.font.color.rgb = CISCO_BLUE if (is_main and i==1) else DARK
        if i == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.add_page_break()


# ── SECTION 1: OVERVIEW ──────────────────────────────────────────────────────

h1(doc, '1  Overview')
add_horizontal_rule(doc)
doc.add_paragraph()

body(doc,
    'The Lab Guide Automator is a Flask-based web application that lets instructors, '
    'technical writers, and lab authors create and maintain professional lab guides '
    'entirely from a browser. It combines structured content editing, screenshot '
    'management, AI-assisted rewriting, and one-click publishing to GitHub Pages — '
    'all in a single tool.'
)

doc.add_paragraph()
h3(doc, 'Key Capabilities')

features = [
    ('Guide Library',        'Create and manage multiple lab guides in one place.'),
    ('Structured Editor',    'Guides are organised into Introduction → Sections → Steps → Conclusion.'),
    ('Drag-to-Reorder',      'Drag sections and content blocks to reorder without losing content.'),
    ('Inline Editing',       'Edit section titles, overviews, step titles, instructions, and body text — all in-place.'),
    ('Screenshot Repository','Central image store with AI-generated captions; attach images to any step or section block.'),
    ('Content Blocks',       'Each section supports an ordered mix of text blocks and screenshot blocks.'),
    ('AI Rewriting',         'One-click AI rewrites for any step, section, introduction, or conclusion.'),
    ('Live Preview',         'Rendered preview with all screenshots in place, per-section step numbering.'),
    ('Multiple Exports',     'Export to Markdown, standalone HTML (Moodle-compatible), or MkDocs site.'),
    ('GitHub Pages Publish', 'Push a Cisco-branded MkDocs site to GitHub Pages with one click.'),
    ('Document Ingest',      'Import an existing .docx file and auto-extract text, headings, and images.'),
]
for label, detail in features:
    bold_bullet(doc, f'{label}:', detail)

doc.add_page_break()


# ── SECTION 2: SYSTEM REQUIREMENTS ───────────────────────────────────────────

h1(doc, '2  System Requirements')
add_horizontal_rule(doc)
doc.add_paragraph()

feature_table(doc, [
    ('Operating System',  'macOS 12+, Linux (Ubuntu 20.04+), Windows 10/11 (WSL2 recommended)'),
    ('Python',            '3.11 or higher'),
    ('Package Manager',   'uv  (recommended)  or  pip + venv'),
    ('Git',               '2.30+  (required for GitHub Pages publishing)'),
    ('Browser',           'Chrome 110+, Firefox 115+, Safari 16+, Edge 110+'),
    ('Disk Space',        '~200 MB for dependencies; additional space for screenshots and exports'),
    ('RAM',               '512 MB minimum; 2 GB recommended when using AI features'),
    ('Internet',          'Required for AI features (GitHub Copilot API) and GitHub Pages publishing'),
], ['Component', 'Requirement'])

note_box(doc,
    'AI features (rewrites, captions) require a valid GitHub Copilot licence associated with '
    'the account whose token is configured in .env. The rest of the tool works fully offline.'
)

doc.add_page_break()


# ── SECTION 3: INSTALLATION ───────────────────────────────────────────────────

h1(doc, '3  Installation')
add_horizontal_rule(doc)
doc.add_paragraph()

# 3.1
h2(doc, '3.1  Prerequisites')
body(doc, 'Before installing, ensure the following are available on your system:')

step_table(doc, [
    ('Install Python 3.11+',
     'Download from https://python.org or use your system package manager.\n'
     'Verify: python3 --version'),
    ('Install uv (recommended package manager)',
     'curl -LsSf https://astral.sh/uv/install.sh | sh\n'
     'Then restart your terminal.'),
    ('Install Git',
     'Download from https://git-scm.com or use your package manager.\n'
     'Verify: git --version'),
    ('Obtain a GitHub Personal Access Token (for publishing)',
     'GitHub → Settings → Developer Settings → Personal Access Tokens → Tokens (classic)\n'
     'Scopes required: repo (full), workflow'),
])

# 3.2
h2(doc, '3.2  Clone the Repository')
body(doc, 'Clone the lab-guide-automator repository to your local machine:')
code_block(doc,
    'git clone https://github.com/mokuma56/lab-guide-automator.git\n'
    'cd lab-guide-automator'
)

note_box(doc,
    'If you are working from a zip archive instead, extract it and open a terminal '
    'in the extracted folder before continuing.'
)

# 3.3
h2(doc, '3.3  Install Dependencies')
body(doc, 'Install all Python dependencies using uv (recommended):')
code_block(doc, 'uv sync')
body(doc, 'Or using pip with a virtual environment:')
code_block(doc,
    'python3 -m venv .venv\n'
    'source .venv/bin/activate        # Windows: .venv\\Scripts\\activate\n'
    'pip install -r requirements.txt'
)

note_box(doc,
    'The first install may take 1–2 minutes while downloading packages. '
    'Subsequent runs use the cached lock file and are near-instant.'
)

# 3.4
h2(doc, '3.4  Environment Configuration')
body(doc,
    'The application reads configuration from a .env file in the project root. '
    'A template is provided — copy it and fill in your values:'
)
code_block(doc, 'cp .env.example .env')
body(doc, 'Open .env in any text editor and set the following values:')

feature_table(doc, [
    ('GITHUB_TOKEN',       'Required for AI features and publishing. Your GitHub Personal Access Token.'),
    ('GITHUB_REPO',        'Optional default repo for MkDocs publishing, e.g. mokuma56/my-lab-guides'),
    ('GITHUB_BRANCH',      'Branch to publish to. Default: main'),
    ('PORT',               'Port the dashboard listens on. Default: 5051'),
    ('DATA_DIR',           'Directory for guide data. Default: ./data'),
], ['Variable', 'Description'])

warning_box(doc,
    'Never commit your .env file to version control. It is already listed in .gitignore. '
    'Keep your GitHub token private.'
)

# 3.5
h2(doc, '3.5  First Launch')
body(doc, 'Start the dashboard with:')
code_block(doc, 'uv run python dashboard.py')
body(doc, 'Or to run in the background (macOS/Linux):')
code_block(doc,
    'nohup uv run python dashboard.py > /tmp/lab_guide_dashboard.log 2>&1 &\n'
    'echo "Dashboard running. Logs: /tmp/lab_guide_dashboard.log"'
)
body(doc, 'Open your browser and navigate to:')
code_block(doc, 'http://localhost:5051')

note_box(doc,
    'To stop the background process: pkill -f "python.*dashboard.py"\n'
    'To view live logs:  tail -f /tmp/lab_guide_dashboard.log'
)

doc.add_page_break()


# ── SECTION 4: DASHBOARD TOUR ─────────────────────────────────────────────────

h1(doc, '4  Dashboard Tour')
add_horizontal_rule(doc)
doc.add_paragraph()

body(doc,
    'The dashboard is a single-page application with two main areas: the Guide Library '
    'panel on the left and the Guide Editor on the right.'
)

h2(doc, '4.1  Guide Library (Left Panel)')
body(doc,
    'The left panel lists all guides stored in data/guides/. From here you can:'
)
bullet(doc, 'Click a guide title to open it in the editor.')
bullet(doc, 'Click + New to create a new blank guide with default sections.')
bullet(doc, 'Delete a guide using the ✕ button next to its name.')

h2(doc, '4.2  Content Tab')
body(doc,
    'The Content tab is the main editing area. It is divided into three layers:'
)
bold_bullet(doc, 'Introduction card:', 'Free-text introduction shown at the top of the guide. Click Edit to modify, or ✦ AI to rewrite with AI.')
bold_bullet(doc, 'Sections area:', 'The ordered list of sections. Each section contains an overview, optional content blocks, and numbered steps.')
bold_bullet(doc, 'Conclusion card:', 'Free-text conclusion at the end of the guide. Click Edit to modify.')
body(doc, '')
body(doc, 'The + Add Section button sits between the Introduction and Sections area for easy access.')

h2(doc, '4.3  Preview Tab')
body(doc,
    'Click the Preview tab to see a fully-rendered view of your guide with all screenshots '
    'in place. Steps are numbered per-section (1, 2, 3… restarting for each section). '
    'Click any screenshot to zoom it to full size.'
)

h2(doc, '4.4  Export Tab')
body(doc, 'From the Export tab you can:')
bullet(doc, 'Download the guide as a Markdown (.md) file.')
bullet(doc, 'Download the guide as a standalone HTML file (Moodle-compatible).')
bullet(doc, 'Configure a GitHub repository and branch for MkDocs publishing.')
bullet(doc, 'Click Publish to GitHub Pages to build a Cisco-branded MkDocs site and push it live.')

h2(doc, '4.5  AI Tools Tab')
body(doc, 'The AI Tools tab provides:')
bullet(doc, 'Bulk AI review — analyse the full guide and receive improvement suggestions.')
bullet(doc, 'One-click application of individual AI suggestions to specific steps or sections.')

doc.add_page_break()


# ── SECTION 5: CREATING YOUR FIRST GUIDE ─────────────────────────────────────

h1(doc, '5  Creating Your First Guide')
add_horizontal_rule(doc)
doc.add_paragraph()

body(doc,
    'Follow these steps to create a new guide from scratch:'
)

step_table(doc, [
    ('Click + New in the Guide Library panel.',
     'A modal dialog appears.'),
    ('Enter the guide title.',
     'For example: "Getting Started with Cisco Cloud Control"'),
    ('Fill in optional fields: Author, Difficulty, Duration.',
     'These appear in the guide metadata and exported documents.'),
    ('Click Create Guide.',
     'The guide opens in the editor with three default sections already created:\n'
     '  • Introduction — with a "Lab Overview" step\n'
     '  • Section 1 — with a placeholder step ready to rename\n'
     '  • Conclusion — with a "Summary" step'),
    ('Edit the Introduction card.',
     'Click Edit in the Introduction card header. Type your introduction text and click Save.'),
    ('Rename Section 1.',
     'Click ✎ Rename in the section header. Type the new name (e.g. "Browse the Integration Catalog") and press Enter.'),
    ('Edit the section Overview.',
     'Click Edit in the overview bar below the section header. Add a one-line description and click Save.'),
    ('Edit the default step.',
     'Click Edit on the step card. Update the Title, Instruction, and Expected Result. Click Save.'),
    ('Add more steps.',
     'Click + Add Step at the bottom of the section to add additional steps.'),
    ('Add more sections.',
     'Click + Add Section to create another section. Drag sections by the ⠿ handle to reorder them.'),
    ('Edit the Conclusion card.',
     'Click Edit in the Conclusion card header and write your closing summary.'),
    ('Preview the guide.',
     'Click the Preview tab to review the fully-rendered guide with step numbers and screenshots.'),
])

doc.add_page_break()


# ── SECTION 6: EDITING GUIDES ─────────────────────────────────────────────────

h1(doc, '6  Editing Guides')
add_horizontal_rule(doc)
doc.add_paragraph()

h2(doc, '6.1  Introduction & Conclusion')
body(doc,
    'Both the Introduction and Conclusion cards support inline editing. '
    'Each card has two buttons in its header:'
)
bold_bullet(doc, 'Edit:', 'Expands a textarea with the current text. Press Save to commit or Cancel to discard.')
bold_bullet(doc, '✦ AI:', 'Opens the AI feedback modal. Add optional guidance and the AI will rewrite the text.')

h2(doc, '6.2  Sections')
body(doc, 'Each section card has the following controls:')

feature_table(doc, [
    ('⠿  (drag handle)',  'Drag the entire section to a new position. Steps are renumbered automatically.'),
    ('Number badge',      'Shows the section\'s position in the guide (1, 2, 3…). Updates on reorder.'),
    ('Section title',     'Displayed in the header.'),
    ('✎ Rename',          'Click to edit the section title inline. Press Enter to save, Escape to cancel.'),
    ('✦ AI',              'AI-rewrite the section overview text.'),
    ('✕',                 'Delete the section and all its steps. Prompts for confirmation.'),
    ('Overview bar',      'One-line description shown under the header. Click Edit to modify.'),
    ('+ Add Step',        'Adds a new step at the bottom of the section.'),
], ['Control', 'Description'])

note_box(doc,
    'Deleting a section is permanent and cannot be undone from the UI. '
    'The guide JSON is saved immediately after deletion.'
)

h2(doc, '6.3  Steps')
body(doc,
    'Steps are the core content unit. Each step card contains:'
)
bold_bullet(doc, 'Step number badge:', 'Blue circle showing the step\'s position within its section (restarts at 1 for each section).')
bold_bullet(doc, 'Step title:', 'Displayed in the card header.')
bold_bullet(doc, 'Edit button:', 'Opens an inline edit panel with three editable fields:')
bullet(doc, 'Title — the step heading.', level=1)
bullet(doc, 'Instruction — the main step text (supports Markdown).', level=1)
bullet(doc, 'Expected Result — what the learner should see when done.', level=1)
bold_bullet(doc, '✦ AI:', 'AI-rewrite the step instruction with optional feedback.')
bold_bullet(doc, 'Screenshots panel:', 'Attach one or more screenshots to the step (see Section 7).')

body(doc, '')
body(doc, 'To edit a step:')
step_table(doc, [
    ('Click Edit on the step card.', 'The edit panel expands below the step header.'),
    ('Modify the Title, Instruction, and/or Expected Result.', ''),
    ('Click Save.', 'Changes are saved to the guide immediately.'),
    ('Click Cancel to discard changes.', ''),
])

h2(doc, '6.4  Content Blocks')
body(doc,
    'Each section supports an ordered list of content blocks between the overview and the steps. '
    'Blocks let you interleave explanatory text and screenshots in any order.'
)
bold_bullet(doc, 'Text block:', 'A freeform markdown text area. Click the text to edit inline.')
bold_bullet(doc, 'Screenshot block:', 'An image with caption. Pick from the repository, upload directly, or request an AI caption.')
body(doc, '')
body(doc, 'To add a block, use the divider row between existing blocks:')
bullet(doc, '+ Text — inserts a new empty text block at that position.')
bullet(doc, '🖼 Screenshot — inserts a new screenshot block and opens the repository picker.')
body(doc, 'Drag the ⠿ handle on any block to reorder it within the section.')

doc.add_page_break()


# ── SECTION 7: SCREENSHOTS ────────────────────────────────────────────────────

h1(doc, '7  Screenshots')
add_horizontal_rule(doc)
doc.add_paragraph()

h2(doc, '7.1  Screenshot Repository')
body(doc,
    'All screenshots are stored in a central repository (data/screenshots/). '
    'Access the repository from the Screenshots tab in the left panel, or via the '
    '"+ Add from Repository" button on any step or section block.'
)
body(doc, 'From the repository you can:')
bullet(doc, 'Browse all screenshots as a thumbnail grid.')
bullet(doc, 'Search by filename or caption.')
bullet(doc, 'Upload new images (PNG, JPG, JPEG).')
bullet(doc, 'Request an AI-generated caption for any image (✦ AI Caption button).')
bullet(doc, 'Edit captions manually by clicking the caption text.')
bullet(doc, 'Delete images from the repository.')

h2(doc, '7.2  Attaching Screenshots to Steps')
body(doc, 'There are two ways to attach a screenshot to a step:')
bold_bullet(doc, 'From the Repository:',
    'Click "+ Add from Repository" in the step\'s screenshot panel. '
    'Browse or search the repository, click an image to select it (blue border appears), '
    'then click "Attach to Step".')
bold_bullet(doc, 'Direct Upload:',
    'Click "⬆ Upload" in the step\'s screenshot panel. '
    'Select one or more image files — they are uploaded to the repository and '
    'automatically attached to the step.')
body(doc, '')
body(doc, 'Once attached, screenshots appear as draggable thumbnails. You can:')
bullet(doc, 'Drag thumbnails to reorder them within the step.')
bullet(doc, 'Click the ✎ button on a thumbnail to edit its caption.')
bullet(doc, 'Click the ✕ button to remove it from the step (does not delete from the repository).')
bullet(doc, 'Click the thumbnail image itself to preview it full-size.')

h2(doc, '7.3  AI Captions')
body(doc,
    'The ✦ AI Caption button (available on both repository images and screenshot blocks) '
    'sends the image to the configured AI vision model and returns a descriptive caption. '
    'The caption is saved to the screenshot metadata and used as the alt text in exports.'
)

note_box(doc,
    'AI captions require a GitHub Copilot account with vision model access. '
    'If the request fails, the existing caption is preserved unchanged.'
)

doc.add_page_break()


# ── SECTION 8: INGESTING EXISTING DOCUMENTS ───────────────────────────────────

h1(doc, '8  Ingesting Existing Documents')
add_horizontal_rule(doc)
doc.add_paragraph()

body(doc,
    'If you have an existing lab guide as a Word (.docx) file, the automator can import '
    'it and convert it into a structured guide automatically.'
)

h2(doc, 'Using the Ingest Feature from the Dashboard')
step_table(doc, [
    ('Open the guide editor for the guide you want to populate.',
     'Or create a new blank guide first.'),
    ('Click the Ingest tab (if available) or use the CLI script.',
     'See the CLI section below for the command-line approach.'),
    ('Select your .docx file.',
     'The ingestor parses headings (H1 → sections, H2/H3 → steps), '
     'body text (→ instructions), and embedded images (→ screenshot repository).'),
    ('Review the imported content.',
     'Sections and steps are created automatically. '
     'Screenshots are extracted and named descriptively based on their context.'),
    ('Clean up as needed.',
     'Rename sections, reorder steps, and attach screenshots to the correct steps '
     'using the standard editor tools.'),
])

h2(doc, 'Using the CLI Script')
body(doc, 'For batch ingest, the reingest_docx.py script is available:')
code_block(doc,
    '# Re-ingest a docx into an existing guide\n'
    'uv run python reingest_docx.py \\\n'
    '  --docx path/to/your-guide.docx \\\n'
    '  --guide-id 84ff424d-ec65-4579-8be3-b9b7101bf6ea'
)
note_box(doc,
    'Ingest replaces the guide\'s sections with the parsed content. '
    'Make a backup of your guide JSON before re-ingesting into an existing guide.'
)

doc.add_page_break()


# ── SECTION 9: EXPORTING & PUBLISHING ─────────────────────────────────────────

h1(doc, '9  Exporting & Publishing')
add_horizontal_rule(doc)
doc.add_paragraph()

h2(doc, '9.1  Markdown Export')
body(doc,
    'Downloads the guide as a single .md file with all text content. '
    'Screenshot paths are relative — suitable for use in any Markdown-based system.'
)
body(doc, 'Click: Export tab → Download Markdown')

h2(doc, '9.2  HTML Export')
body(doc,
    'Downloads a self-contained HTML file with embedded Cisco-brand CSS. '
    'Compatible with Moodle and other LMS platforms that accept HTML uploads.'
)
body(doc, 'Click: Export tab → Download HTML')

h2(doc, '9.3  MkDocs / GitHub Pages Publish')
body(doc,
    'Builds a full Cisco-branded MkDocs documentation site and pushes it to GitHub Pages. '
    'The published site includes:'
)
bullet(doc, 'Navigation sidebar with all sections and steps.')
bullet(doc, 'Lightbox-enabled screenshot galleries.')
bullet(doc, 'Cisco color palette and CiscoSansTT-inspired typography.')
bullet(doc, 'A custom landing page with the lab title and "Get Started" button.')

body(doc, '')
body(doc, 'Setup (one time per guide):')
step_table(doc, [
    ('Open the Export tab.', ''),
    ('Enter the GitHub repository URL.',
     'Example: https://github.com/youruser/your-lab-guide.git'),
    ('Set the branch (default: main).', ''),
    ('Click Save Config.', 'Settings are stored in the guide JSON.'),
    ('Click Publish to GitHub Pages.',
     'A live log stream shows the MkDocs build and git push progress. '
     'The site is live at https://youruser.github.io/your-lab-guide/ within ~60 seconds.'),
])

warning_box(doc,
    'The GitHub repository must exist before publishing. Create it at github.com first. '
    'GitHub Pages must be enabled in the repository Settings → Pages → Source: gh-pages branch.'
)

doc.add_page_break()


# ── SECTION 10: AI FEATURES ───────────────────────────────────────────────────

h1(doc, '10  AI Features')
add_horizontal_rule(doc)
doc.add_paragraph()

body(doc,
    'The Lab Guide Automator integrates with the GitHub Copilot API to provide '
    'AI-assisted content generation throughout the authoring workflow.'
)

feature_table(doc, [
    ('Rewrite Step',
     'Click ✦ AI on any step card. Optionally add feedback (e.g. "make it more specific, add CLI commands"). The AI rewrites the instruction and expected result.'),
    ('Rewrite Section Overview',
     'Click ✦ AI in the section header. The AI rewrites the section overview based on the section\'s existing steps.'),
    ('Rewrite Introduction',
     'Click ✦ AI in the Introduction card header. AI generates a new introduction based on the guide title, metadata, and sections.'),
    ('Rewrite Conclusion',
     'Click ✦ AI in the Conclusion card header. AI summarises what the learner accomplished.'),
    ('AI Caption (Screenshot)',
     'Click ✦ AI Caption on any screenshot in the repository or a screenshot block. The vision model describes the image content.'),
    ('Bulk Review (AI Tools tab)',
     'Analyses the entire guide and returns a list of improvement suggestions. Each suggestion can be applied individually with one click.'),
], ['Feature', 'How to Use'])

note_box(doc,
    'AI requests are routed through the GitHub Copilot token configured in .env. '
    'Requests are not logged or stored by the application. '
    'Network connectivity to api.githubcopilot.com is required.'
)

doc.add_page_break()


# ── SECTION 11: KEYBOARD & INTERACTION REFERENCE ──────────────────────────────

h1(doc, '11  Keyboard & Interaction Reference')
add_horizontal_rule(doc)
doc.add_paragraph()

feature_table(doc, [
    ('Enter',           'Confirm / save in any inline text input (section title, step title).'),
    ('Escape',          'Cancel / discard in any inline text input.'),
    ('Click & drag ⠿', 'Reorder sections, content blocks, and screenshot thumbnails.'),
    ('Double-click',    'Not required — all editing is via explicit Edit buttons.'),
    ('Click thumbnail', 'Preview screenshot full-size in a modal overlay.'),
    ('Click modal bg',  'Close any open modal.'),
], ['Interaction', 'Action'])

doc.add_page_break()


# ── SECTION 12: TROUBLESHOOTING ───────────────────────────────────────────────

h1(doc, '12  Troubleshooting')
add_horizontal_rule(doc)
doc.add_paragraph()

issues = [
    (
        'Dashboard will not start / "Address already in use"',
        'Another process is using port 5051. Kill it and restart:\n'
        '  pkill -f "python.*dashboard.py"\n'
        '  uv run python dashboard.py'
    ),
    (
        'Changes are not saved after editing',
        'The dashboard caches guides in memory. If you edited guide JSON files directly on disk, '
        'restart the dashboard to reload from disk. Always use the UI for edits.'
    ),
    (
        'AI features return errors',
        'Check that GITHUB_TOKEN in .env is valid and has Copilot access. '
        'Test with: curl -H "Authorization: token YOUR_TOKEN" https://api.github.com/user'
    ),
    (
        'Screenshots not showing in preview',
        'Ensure the screenshot filename matches what is stored in data/screenshots/. '
        'The preview uses /api/screenshots/file/<name> — if the file is missing from disk, the image is hidden automatically.'
    ),
    (
        'MkDocs publish fails',
        'Verify: (1) GitHub repo exists, (2) PAT has repo + workflow scopes, '
        '(3) mkdocs and mkdocs-material are installed (uv sync). '
        'Check the live log stream in the Export tab for the specific error.'
    ),
    (
        'Import (.docx) produces empty sections',
        'The ingestor relies on Word heading styles (Heading 1, Heading 2, etc.). '
        'If your document uses manually formatted bold text instead of heading styles, '
        'apply proper heading styles in Word before importing.'
    ),
    (
        'Port 5051 not accessible on remote machine',
        'By default Flask binds to 127.0.0.1. To allow remote access, edit dashboard.py and '
        'change app.run(host="127.0.0.1") to app.run(host="0.0.0.0"). '
        'Ensure your firewall allows port 5051.'
    ),
]

for title, detail in issues:
    h3(doc, title)
    body(doc, detail)
    doc.add_paragraph()

doc.add_page_break()


# ── SECTION 13: API REFERENCE ─────────────────────────────────────────────────

h1(doc, '13  API Reference')
add_horizontal_rule(doc)
doc.add_paragraph()

body(doc,
    'The dashboard exposes a REST API under /api/. All requests and responses use JSON '
    'unless otherwise noted. The base URL is http://localhost:5051.'
)

h2(doc, 'Guides')
feature_table(doc, [
    ('GET  /api/guides',                        'List all guides (id, title, updated_at).'),
    ('POST /api/guides',                        'Create a new guide. Body: {title, author, difficulty, duration}.'),
    ('GET  /api/guides/<id>',                   'Get full guide JSON.'),
    ('DELETE /api/guides/<id>',                 'Delete a guide.'),
    ('POST /api/guides/<id>/metadata',          'Update guide metadata fields.'),
    ('GET  /api/guides/<id>/preview',           'Get rendered HTML preview (with live screenshot URLs).'),
    ('POST /api/guides/<id>/introduction',      'Save introduction text. Body: {text}.'),
    ('POST /api/guides/<id>/conclusion',        'Save conclusion text. Body: {text}.'),
    ('POST /api/guides/<id>/sections/reorder',  'Reorder sections. Body: {order: [sec_id, ...]}.'),
], ['Endpoint', 'Description'])

h2(doc, 'Sections')
feature_table(doc, [
    ('POST /api/guides/<id>/section',                   'Add a section. Body: {title, overview}.'),
    ('POST /api/guides/<id>/section/<sec_id>',          'Update section fields. Body: {title?, overview?}.'),
    ('DELETE /api/guides/<id>/section/<sec_id>',        'Delete section and its steps.'),
    ('GET  /api/guides/<id>/section/<sec_id>/blocks',   'List content blocks.'),
    ('POST /api/guides/<id>/section/<sec_id>/blocks',   'Add a block. Body: {type, content?, path?, caption?}.'),
    ('POST .../blocks/reorder',                         'Reorder blocks. Body: {order: [block_id, ...]}.'),
], ['Endpoint', 'Description'])

h2(doc, 'Steps')
feature_table(doc, [
    ('POST /api/guides/<id>/step',                              'Add a step. Body: {section_id, title, description}.'),
    ('POST /api/guides/<id>/step/<step_id>',                    'Update step. Body: {title?, instruction?, expected_result?, notes?}.'),
    ('POST /api/guides/<id>/step/<step_id>/screenshots',        'Attach a screenshot. Body: {path, caption?}.'),
    ('DELETE /api/guides/<id>/step/<step_id>/screenshots/<idx>','Remove screenshot at index.'),
], ['Endpoint', 'Description'])

h2(doc, 'Screenshots')
feature_table(doc, [
    ('GET  /api/screenshots',               'List all screenshots with captions and metadata.'),
    ('GET  /api/screenshots/file/<name>',   'Serve an image file.'),
    ('POST /api/screenshots/upload',        'Upload image(s). Multipart form with file field.'),
    ('POST /api/screenshots/<name>/caption','Update caption. Body: {caption} or {ai: true}.'),
    ('POST /api/screenshots/<name>/delete', 'Delete image from repository.'),
], ['Endpoint', 'Description'])

h2(doc, 'Export & Publish')
feature_table(doc, [
    ('GET  /api/guides/<id>/export/markdown', 'Download guide as .md file.'),
    ('GET  /api/guides/<id>/export/html',     'Download guide as standalone .html file.'),
    ('POST /api/guides/<id>/sync-config',     'Save GitHub repo/branch config. Body: {github_repo, github_branch}.'),
    ('GET  /api/guides/<id>/publish',         'SSE stream — builds MkDocs site and pushes to GitHub Pages.'),
], ['Endpoint', 'Description'])


# ── BACK COVER ────────────────────────────────────────────────────────────────

doc.add_page_break()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

cover2 = doc.add_table(rows=1, cols=1)
cover2.style = 'Table Grid'
c2 = cover2.cell(0,0)
set_cell_bg(c2, '005073')
p2 = c2.paragraphs[0]
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run('\n\nLab Guide Automator\n')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = WHITE
r2 = p2.add_run('Version 1.0  ·  July 2026\n')
r2.font.size = Pt(12)
r2.font.color.rgb = CISCO_CYAN
r3 = p2.add_run('\nhttps://github.com/mokuma56/lab-guide-automator\n\n')
r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0xCC, 0xE8, 0xF5)


# ── SAVE ──────────────────────────────────────────────────────────────────────

out_path = Path(__file__).parent / 'Lab_Guide_Automator_User_Guide.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
