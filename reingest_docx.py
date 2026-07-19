"""
Re-ingest agent-builder-customer-guide.docx into the existing guide JSON,
extracting all images and associating them with the correct steps.

Run from project root:
    uv run python reingest_docx.py
"""
from __future__ import annotations
import json, re, shutil, uuid
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

DOCX_PATH = Path("data/uploads/agent-builder-customer-guide.docx")
GUIDE_ID  = "84ff424d-ec65-4579-8be3-b9b7101bf6ea"
GUIDE_PATH = Path(f"data/guides/{GUIDE_ID}.json")
SCREENSHOTS_DIR = Path("data/screenshots")
SCREENSHOTS_DIR.mkdir(parents=True, exist_ok=True)

# ──────────────────────────────────────────────────────────────
# 1.  Extract all images from the docx (keyed by relationship ID)
# ──────────────────────────────────────────────────────────────
doc = Document(str(DOCX_PATH))

img_bytes: dict[str, tuple[bytes, str]] = {}   # rId → (data, ext)
for rId, rel in doc.part.rels.items():
    if "image" in rel.reltype:
        part = rel.target_part
        ext  = part.content_type.split("/")[-1].replace("jpeg", "jpg")
        if ext not in ("png", "jpg", "jpeg", "gif", "svg"):
            ext = "png"
        img_bytes[rId] = (part.blob, ext)

print(f"Found {len(img_bytes)} images in docx")

# ──────────────────────────────────────────────────────────────
# 2.  Walk paragraphs, collect images with their captions and
#     the nearest preceding heading/text for context
# ──────────────────────────────────────────────────────────────

def para_images(para) -> list[str]:
    """Return list of rIds for images embedded in this paragraph."""
    blips = para._element.findall(".//" + qn("a:blip"))
    return [b.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            for b in blips]

# (para_idx, rId, caption_text)
img_records: list[tuple[int, str, str]] = []

for i, para in enumerate(doc.paragraphs):
    rids = para_images(para)
    if rids:
        # caption is usually the next paragraph
        caption = ""
        if i + 1 < len(doc.paragraphs):
            nxt = doc.paragraphs[i + 1]
            if "Caption" in nxt.style.name or nxt.style.name in ("Image Caption",):
                caption = nxt.text.strip()
        for rid in rids:
            img_records.append((i, rid, caption))

print(f"Mapped {len(img_records)} image placements")

# ──────────────────────────────────────────────────────────────
# 3.  Save images to data/screenshots/
# ──────────────────────────────────────────────────────────────
saved_images: list[dict] = []   # {para_idx, filename, caption}
for para_idx, rid, caption in img_records:
    data, ext = img_bytes.get(rid, (b"", "png"))
    if not data:
        continue
    slug = re.sub(r"[^\w]+", "-", caption[:60].lower()).strip("-") or rid
    fname = f"{GUIDE_ID[:8]}-{slug}.{ext}"
    dest  = SCREENSHOTS_DIR / fname
    dest.write_bytes(data)
    saved_images.append({"para_idx": para_idx, "path": f"screenshots/{fname}", "caption": caption})
    print(f"  saved {fname}")

# ──────────────────────────────────────────────────────────────
# 4.  Re-parse docx paragraphs into a structured guide
#     (exact text, proper section/step hierarchy)
# ──────────────────────────────────────────────────────────────

def slugify(t): return re.sub(r"[^\w]+", "-", t.lower()).strip("-")

# Map para_idx → image list
from collections import defaultdict
para_to_images = defaultdict(list)
for rec in saved_images:
    para_to_images[rec["para_idx"]].append(rec)

paras = doc.paragraphs

# Build section/step structure by walking paragraphs
from dataclasses import dataclass, field

@dataclass
class Step:
    order: int
    title: str
    lines: list[str] = field(default_factory=list)
    screenshots: list[dict] = field(default_factory=list)
    code_blocks: list[str] = field(default_factory=list)
    expected_result: str = ""
    notes: str = ""

@dataclass
class Section:
    title: str
    overview_lines: list[str] = field(default_factory=list)
    steps: list[Step] = field(default_factory=list)

sections: list[Section] = []
cur_section: Section | None = None
cur_step: Step | None = None
intro_lines: list[str] = []
step_counter = 0

def flush_step():
    if cur_step and cur_section:
        cur_section.steps.append(cur_step)

def flush_section():
    if cur_section:
        flush_step()
        sections.append(cur_section)

TEXT_CAPTION_STYLES = {"Image Caption", "Caption"}  # text-only caption lines to skip

for i, para in enumerate(paras):
    text  = para.text.strip()
    style = para.style.name
    imgs  = para_to_images.get(i, [])
    has_img = bool(para_images(para))

    # Skip empty paragraphs with no images
    if not text and not has_img:
        continue
    # Skip text-only caption lines (already captured in the image records)
    if not has_img and (style in TEXT_CAPTION_STYLES or style == "Image Caption"):
        continue

    # ── Headings ────────────────────────────────────────────────
    if style == "Heading 1" or style == "Title":
        # document title — skip, captured in metadata
        continue

    elif style == "Heading 2":
        flush_section()
        cur_section = Section(title=text)
        cur_step = None
        step_counter = 0

    elif style in ("Heading 3", "Heading 4"):
        flush_step()
        step_counter += 1
        cur_step = Step(order=step_counter, title=text)

    else:
        # Body text / Normal / Compact / Block Text etc.
        if text:
            if cur_step is not None:
                cur_step.lines.append(text)
            elif cur_section is not None:
                cur_section.overview_lines.append(text)
            else:
                intro_lines.append(text)

    # Attach images at this para position to wherever we currently are
    for img in imgs:
        if cur_step is not None:
            cur_step.screenshots.append(img)
        elif cur_section is not None:
            # No sub-step yet — create a synthetic "Overview" step to hold images
            if not cur_section.steps:
                step_counter += 1
                cur_step = Step(order=step_counter, title="Overview",
                                lines=list(cur_section.overview_lines))
                cur_section.overview_lines = []
            cur_step.screenshots.append(img)

flush_section()

print(f"\nParsed {len(sections)} sections:")
for s in sections:
    shots = sum(len(st.screenshots) for st in s.steps)
    print(f"  {s.title}: {len(s.steps)} steps, {shots} screenshots")

# ──────────────────────────────────────────────────────────────
# 5.  Load existing guide JSON and replace content
# ──────────────────────────────────────────────────────────────
with open(GUIDE_PATH) as f:
    guide = json.load(f)

def make_step(step: Step) -> dict:
    instruction = "\n\n".join(step.lines)
    return {
        "id": str(uuid.uuid4())[:8],
        "order": step.order,
        "title": step.title,
        "instruction": instruction,
        "expected_result": step.expected_result,
        "notes": step.notes,
        "screenshots": [
            {"path": img["path"], "timestamp_s": 0.0, "caption": img["caption"]}
            for img in step.screenshots
        ],
        "code_blocks": step.code_blocks,
        "verified": False,
    }

def make_section(sec: Section) -> dict:
    return {
        "id": str(uuid.uuid4())[:8],
        "title": sec.title,
        "overview": "\n\n".join(sec.overview_lines),
        "steps": [make_step(st) for st in sec.steps],
    }

guide["introduction"] = "\n\n".join(intro_lines)
guide["sections"] = [make_section(s) for s in sections]

from datetime import datetime
guide["updated_at"] = datetime.utcnow().isoformat()

with open(GUIDE_PATH, "w") as f:
    json.dump(guide, f, indent=2)

total_shots = sum(
    len(st["screenshots"])
    for sec in guide["sections"]
    for st in sec["steps"]
)
print(f"\nGuide saved. {len(guide['sections'])} sections, {total_shots} total screenshots.")
print(f"Path: {GUIDE_PATH}")
